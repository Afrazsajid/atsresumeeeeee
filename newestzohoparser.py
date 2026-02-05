import os
import io
import re
import json
import math
import time
import logging
import traceback
from dataclasses import dataclass, field
from pathlib import Path
from datetime import datetime, date
from typing import Dict, List, Optional, Tuple, Any
from urllib.parse import urlparse, parse_qs

import streamlit as st
import pandas as pd
import numpy as np
import requests

# Optional dependencies
try:
    import pdfplumber
    PDF_AVAILABLE = True
except Exception:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    import spacy
    SPACY_AVAILABLE = True
except Exception:
    SPACY_AVAILABLE = False

try:
    import phonenumbers
    PHONE_AVAILABLE = True
except Exception:
    PHONE_AVAILABLE = False

try:
    from email_validator import validate_email, EmailNotValidError
    EMAIL_AVAILABLE = True
except Exception:
    EMAIL_AVAILABLE = False

try:
    from rapidfuzz import fuzz
    RAPIDFUZZ_AVAILABLE = True
except Exception:
    RAPIDFUZZ_AVAILABLE = False

# OCR optional
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False


# ---------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------
logger = logging.getLogger("cv_parser")
logger.setLevel(logging.INFO)
_handler = logging.StreamHandler()
_handler.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
if not logger.handlers:
    logger.addHandler(_handler)


# ---------------------------------------------------------------------
# Resume filename pattern
# Matches stems that START with Resume or CV (case-insensitive)
# Examples that match:
#   Resume-Manisha_mukeshbhai-Radadiya.docx
#   Resume - John Smith.pdf
#   CV_Ali-Khan.docx
# ---------------------------------------------------------------------
RESUME_STEM_RE = re.compile(r"^(resume|cv)(?:[-_ ]|$)", re.IGNORECASE)


def is_resume_filename(name: str, supported_exts: Tuple[str, ...], exact_filename: str = "") -> bool:
    if not name:
        return False

    if exact_filename and name.strip().lower() != exact_filename.strip().lower():
        return False

    p = Path(name)
    if p.suffix.lower() not in supported_exts:
        return False

    return bool(RESUME_STEM_RE.match(p.stem))


# ---------------------------------------------------------------------
# Zoho WorkDrive Client
# ---------------------------------------------------------------------
class ZohoWorkDriveClient:
    """
    WorkDrive recursive folder walker with strict resume file filtering.

    Required secrets in .streamlit/secrets.toml:

    ZOHO_CLIENT_ID = "..."
    ZOHO_CLIENT_SECRET = "..."
    ZOHO_REFRESH_TOKEN = "..."

    Optional region configuration:
    ZOHO_ACCOUNTS_DOMAIN = "https://accounts.zoho.com"   # or accounts.zoho.eu, accounts.zoho.in, etc
    ZOHO_API_BASE = "https://www.zohoapis.com"           # or www.zohoapis.eu, www.zohoapis.in, etc
    ZOHO_WORKDRIVE_WEB_BASE = "https://workdrive.zoho.com"
    """

    def __init__(self):
        self.client_id = st.secrets.get("ZOHO_CLIENT_ID", "")
        self.client_secret = st.secrets.get("ZOHO_CLIENT_SECRET", "")
        self.refresh_token = st.secrets.get("ZOHO_REFRESH_TOKEN", "")

        self.accounts_domain = st.secrets.get("ZOHO_ACCOUNTS_DOMAIN", "https://accounts.zoho.com").rstrip("/")
        self.api_base = st.secrets.get("ZOHO_API_BASE", "https://www.zohoapis.com").rstrip("/")
        self.workdrive_web_base = st.secrets.get("ZOHO_WORKDRIVE_WEB_BASE", "https://workdrive.zoho.com").rstrip("/")

        if not (self.client_id and self.client_secret and self.refresh_token):
            raise ValueError("Missing Zoho secrets. Add ZOHO_CLIENT_ID, ZOHO_CLIENT_SECRET, ZOHO_REFRESH_TOKEN to secrets.toml")

        self.workdrive_api_base = f"{self.api_base}/workdrive/api/v1"

    def _get_cached_token(self) -> Optional[str]:
        tok = st.session_state.get("zoho_access_token")
        exp = st.session_state.get("zoho_access_token_expires_at", 0)
        if tok and time.time() < exp - 30:
            return tok
        return None

    def _cache_token(self, token: str, expires_in_sec: int):
        st.session_state["zoho_access_token"] = token
        st.session_state["zoho_access_token_expires_at"] = time.time() + int(expires_in_sec)

    def refresh_access_token(self) -> str:
        cached = self._get_cached_token()
        if cached:
            return cached

        url = f"{self.accounts_domain}/oauth/v2/token"
        data = {
            "grant_type": "refresh_token",
            "client_id": self.client_id,
            "client_secret": self.client_secret,
            "refresh_token": self.refresh_token,
        }
        r = requests.post(url, data=data, timeout=30)
        if r.status_code != 200:
            raise RuntimeError(f"Token refresh failed ({r.status_code}). {r.text}")

        j = r.json()
        token = j.get("access_token")
        expires = int(j.get("expires_in_sec") or j.get("expires_in") or 3600)
        if not token:
            raise RuntimeError(f"Token refresh response missing access_token. {j}")

        self._cache_token(token, expires)
        return token

    def _headers(self) -> Dict[str, str]:
        return {
            "Authorization": f"Zoho-oauthtoken {self.refresh_access_token()}",
            "Accept": "application/vnd.api+json",
        }

    def _request(self, method: str, url: str, **kwargs) -> requests.Response:
        r = requests.request(method, url, headers=self._headers(), timeout=kwargs.pop("timeout", 60), **kwargs)
        if r.status_code == 401:
            st.session_state.pop("zoho_access_token", None)
            st.session_state.pop("zoho_access_token_expires_at", None)
            r = requests.request(method, url, headers=self._headers(), timeout=kwargs.pop("timeout", 60), **kwargs)
        return r

    @staticmethod
    def extract_resource_id(folder_link_or_id: str) -> str:
        """
        Tries hard to extract a WorkDrive folder ID from common link formats.
        If it cannot, it expects user to paste the folder ID directly.
        """
        s = (folder_link_or_id or "").strip()
        if not s:
            raise ValueError("Folder link or id is empty.")

        # If user already pasted an ID
        if re.fullmatch(r"[A-Za-z0-9]{8,}", s):
            return s

        # Parse URL path and query
        u = urlparse(s)
        parts = [p for p in u.path.split("/") if p]
        # Look for any token that resembles an ID
        for p in reversed(parts):
            if re.fullmatch(r"[A-Za-z0-9]{8,}", p):
                return p

        qs = parse_qs(u.query or "")
        for key in ["id", "res_id", "folder_id", "resource_id"]:
            if key in qs and qs[key]:
                val = qs[key][0]
                if re.fullmatch(r"[A-Za-z0-9]{8,}", val):
                    return val

        m = re.search(r"(?:folder|folders|file|files|id|res_id)=([A-Za-z0-9]{8,})", s, flags=re.I)
        if m:
            return m.group(1)

        raise ValueError("Could not detect folder id from link. Paste the folder id directly.")

    @staticmethod
    def _item_name(it: Dict[str, Any]) -> str:
        attrs = it.get("attributes") or {}
        return (attrs.get("name") or it.get("name") or "").strip()

    @staticmethod
    def _item_id(it: Dict[str, Any]) -> str:
        return (it.get("id") or "").strip()

    @staticmethod
    def _is_folder_item(it: Dict[str, Any]) -> bool:
        t = (it.get("type") or "").lower()
        attrs = it.get("attributes") or {}
        # Different API shapes observed in the wild
        if t == "folders":
            return True
        if attrs.get("is_folder") is True:
            return True
        if (attrs.get("type") or "").lower() == "folder":
            return True
        if (attrs.get("mime_type") or "").lower() in {"folder", "application/vnd.zoho.folder"}:
            return True
        if (attrs.get("file_type") or "").lower() == "folder":
            return True
        # If API returns everything as type "files", folder can still be distinguished by a flag
        return False

    @staticmethod
    def _is_file_item(it: Dict[str, Any]) -> bool:
        # Most APIs mark actual files as type "files"
        # Some return type "files" for both and use flags
        if ZohoWorkDriveClient._is_folder_item(it):
            return False
        t = (it.get("type") or "").lower()
        if t == "files":
            return True
        attrs = it.get("attributes") or {}
        if attrs.get("is_folder") is False:
            return True
        return bool(ZohoWorkDriveClient._item_name(it))

    def _list_children_one_page(self, folder_id: str, offset: int, limit: int) -> Tuple[List[Dict[str, Any]], bool]:
        """
        Tries multiple endpoints because WorkDrive API shapes vary by resource type and tenant.
        Returns (items, has_more).
        """
        endpoints = [
            f"{self.workdrive_api_base}/files/{folder_id}/files",
            f"{self.workdrive_api_base}/folders/{folder_id}/files",
            f"{self.workdrive_api_base}/folders/{folder_id}/children",
            f"{self.workdrive_api_base}/files/{folder_id}/children",
        ]

        last_err = None
        for url in endpoints:
            params = {
                "filter[type]": "all",
                "page[limit]": limit,
                "page[offset]": offset,
            }
            try:
                r = self._request("GET", url, params=params, timeout=60)
                if r.status_code != 200:
                    last_err = f"{url} -> {r.status_code} {r.text[:400]}"
                    continue

                payload = r.json() if r.text else {}
                items = payload.get("data", []) or []

                # Sometimes folder children come in "included"
                included = payload.get("included", []) or []
                if included:
                    items = list(items) + list(included)

                # Keep only dict items
                items = [x for x in items if isinstance(x, dict)]

                has_more = False
                links = payload.get("links") or {}
                if isinstance(links, dict) and links.get("next"):
                    has_more = True
                else:
                    # Offset paging heuristic
                    has_more = len(payload.get("data", []) or []) >= limit

                return items, has_more
            except Exception as e:
                last_err = f"{url} -> {e}"
                continue

        raise RuntimeError(f"Could not list folder children. Last error: {last_err}")

    def list_resume_files_recursive(
        self,
        root_folder_id: str,
        include_subfolders: bool,
        supported_exts: Tuple[str, ...],
        exact_filename: str = "",
        max_items_safety: int = 20000
    ) -> List[Dict[str, Any]]:
        """
        Walks every subfolder (folder of folder of folder) and returns only resume files.
        Output items contain: id, name, virtual_path
        """
        results: List[Dict[str, Any]] = []
        visited = set()

        queue: List[Tuple[str, str]] = [(root_folder_id, "")]  # (folder_id, path_prefix)
        scanned = 0

        while queue:
            folder_id, prefix = queue.pop(0)
            if folder_id in visited:
                continue
            visited.add(folder_id)

            offset = 0
            limit = 50
            loops = 0

            while True:
                loops += 1
                if loops > 5000:
                    raise RuntimeError("Safety stop. Too many pagination loops while listing WorkDrive folder contents.")

                items, has_more = self._list_children_one_page(folder_id, offset=offset, limit=limit)
                scanned += len(items)
                if scanned > max_items_safety:
                    raise RuntimeError("Safety stop. Too many items scanned. Narrow folder or increase max_items_safety.")

                for it in items:
                    name = self._item_name(it)
                    it_id = self._item_id(it)
                    if not it_id or not name:
                        continue

                    if self._is_folder_item(it):
                        if include_subfolders:
                            new_prefix = f"{prefix}{name}/"
                            queue.append((it_id, new_prefix))
                        continue

                    if self._is_file_item(it):
                        if is_resume_filename(name, supported_exts=supported_exts, exact_filename=exact_filename):
                            results.append({
                                "id": it_id,
                                "name": name,
                                "virtual_path": f"{prefix}{name}".strip(),
                            })

                if not has_more:
                    break
                offset += limit

        results.sort(key=lambda x: (x.get("virtual_path") or x.get("name") or "").lower())
        return results

    def download_file_bytes(self, file_id: str) -> bytes:
        """
        Tries multiple download endpoints.
        """
        # 1) WorkDrive web download
        url1 = f"{self.workdrive_web_base}/api/v1/download/{file_id}"
        r = self._request("GET", url1, timeout=120, allow_redirects=True)
        if r.status_code == 200 and r.content:
            return r.content

        # 2) download.zoho.com fallback
        url2 = "https://download.zoho.com/v1/workdrive/download/" + file_id
        r2 = self._request("GET", url2, timeout=120, allow_redirects=True)
        if r2.status_code == 200 and r2.content:
            return r2.content

        # 3) WorkDrive API direct download, some tenants support this
        url3 = f"{self.workdrive_api_base}/download/{file_id}"
        r3 = self._request("GET", url3, timeout=120, allow_redirects=True)
        if r3.status_code == 200 and r3.content:
            return r3.content

        raise RuntimeError(
            f"Download failed for file_id={file_id}. "
            f"status1={r.status_code}, status2={r2.status_code}, status3={r3.status_code}"
        )


# ---------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------
MONTHS = {
    "jan": 1, "january": 1,
    "feb": 2, "february": 2,
    "mar": 3, "march": 3,
    "apr": 4, "april": 4,
    "may": 5,
    "jun": 6, "june": 6,
    "jul": 7, "july": 7,
    "aug": 8, "august": 8,
    "sep": 9, "sept": 9, "september": 9,
    "oct": 10, "october": 10,
    "nov": 11, "november": 11,
    "dec": 12, "december": 12,
}

@dataclass
class AppConfig:
    supported_exts: Tuple[str, ...] = (".pdf", ".docx")
    max_file_mb: int = 15

    spacy_model_preference: Tuple[str, ...] = ("en_core_web_lg", "en_core_web_md", "en_core_web_sm")

    # PDF extraction tuning
    min_extracted_alpha_chars_for_pdf: int = 400
    drop_repeated_header_footer_ratio: float = 0.60

    # OCR
    enable_ocr_fallback: bool = True
    ocr_lang: str = "eng"

    # Skills
    embedding_similarity_threshold: float = 0.70
    keep_unmatched_skill_phrases: bool = True

    # Performance
    max_workers: int = 4

    # Sections
    section_aliases: Dict[str, List[str]] = field(default_factory=lambda: {
        "summary": ["summary", "profile", "professional summary", "about", "objective"],
        "skills": ["skills", "technical skills", "core skills", "key skills", "competencies", "expertise"],
        "experience": ["experience", "work experience", "employment", "career history", "work history", "professional experience"],
        "education": ["education", "qualifications", "academic background"],
        "certifications": ["certifications", "certificates", "licences", "licenses", "training", "accreditations"],
        "contact": ["contact", "contact details", "personal details", "personal information"],
    })

    # Industry tagging
    industry_keywords: Dict[str, List[str]] = field(default_factory=lambda: {
        "security": [
            "sia", "door supervisor", "security officer", "cctv", "steward", "event security", "patrol",
            "access control", "crowd", "incident", "search procedure", "id check", "safeguarding"
        ],
        "warehouse": [
            "picker", "packer", "picking", "packing", "rf scanner", "scan gun", "forklift", "reach truck",
            "loading", "unloading", "inventory", "stock", "dispatch", "goods in", "goods out"
        ],
        "hospitality": [
            "housekeeping", "room attendant", "waiter", "waitress", "bar", "bartender", "front of house",
            "concierge", "porter", "linen porter", "service standards"
        ]
    })

    # Regex patterns
    email_re: re.Pattern = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", re.I)
    uk_postcode_re: re.Pattern = re.compile(r"\b([A-Z]{1,2}\d{1,2}[A-Z]?)\s*(\d[A-Z]{2})\b", re.I)

    date_range_re: re.Pattern = re.compile(
        r"(?P<start>(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
        r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t)?(?:ember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}|\b(19|20)\d{2}\b)"
        r"\s*(?:-+|–|to)\s*"
        r"(?P<end>(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
        r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t)?(?:ember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}|\b(19|20)\d{2}\b|present|current|now)"
        , re.I
    )

    bullet_re: re.Pattern = re.compile(r"^\s*([•\-\*\u2022\u25CF\u25CB\u25AA\u00B7]+|\d+\.)\s+")


CONFIG = AppConfig()


# ---------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------
def dedupe_preserve_order(items: List[str]) -> List[str]:
    seen = set()
    out = []
    for x in items:
        k = x.strip()
        if not k:
            continue
        if k.lower() in seen:
            continue
        seen.add(k.lower())
        out.append(k)
    return out


def normalise_whitespace_keep_lines(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = []
    for ln in text.split("\n"):
        ln2 = re.sub(r"[ \t]+", " ", ln).strip()
        lines.append(ln2)
    out = "\n".join(lines)
    out = re.sub(r"\n{3,}", "\n\n", out).strip()
    return out


def safe_clip01(x: float) -> float:
    return float(max(0.0, min(1.0, x)))


def cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
    if a is None or b is None:
        return 0.0
    na = np.linalg.norm(a)
    nb = np.linalg.norm(b)
    if na == 0.0 or nb == 0.0:
        return 0.0
    return float(np.dot(a, b) / (na * nb))


def looks_like_heading(line: str) -> bool:
    if not line:
        return False
    l = line.strip()
    if len(l) < 3:
        return False
    words = l.split()
    if len(words) > 6:
        return False
    if l.endswith(":"):
        return True
    if l.upper() == l and any(c.isalpha() for c in l) and len(l) <= 40:
        return True
    return False


def normalise_postcode(m1: str, m2: str) -> str:
    return f"{m1.upper()} {m2.upper()}"


def parse_month_year(s: str) -> Tuple[Optional[date], float]:
    s0 = s.strip().lower()
    if re.fullmatch(r"(19|20)\d{2}", s0):
        try:
            y = int(s0)
            return date(y, 1, 1), 0.70
        except Exception:
            return None, 0.0

    m = re.match(r"^(?P<mon>[a-z]{3,9})\s+(?P<yr>(19|20)\d{2})$", s0)
    if not m:
        return None, 0.0
    mon = m.group("mon")
    yr = int(m.group("yr"))
    mon_num = MONTHS.get(mon[:3], None) if mon not in MONTHS else MONTHS[mon]
    if not mon_num:
        mon_num = MONTHS.get(mon, None)
    if not mon_num:
        return None, 0.0
    return date(yr, mon_num, 1), 0.90


def date_to_yyyy_mm(d: date) -> str:
    return f"{d.year:04d}-{d.month:02d}"


def months_between(a: date, b: date) -> int:
    return (b.year - a.year) * 12 + (b.month - a.month)


# ---------------------------------------------------------------------
# spaCy loader
# ---------------------------------------------------------------------
@st.cache_resource(show_spinner=False)
def load_spacy_model():
    if not SPACY_AVAILABLE:
        return None
    for name in CONFIG.spacy_model_preference:
        try:
            nlp = spacy.load(name)
            return nlp
        except Exception:
            continue
    return None


# ---------------------------------------------------------------------
# Skill catalogue
# ---------------------------------------------------------------------
DEFAULT_SKILL_CATALOGUE: Dict[str, Any] = {
    "canonical": [
        "SIA",
        "SIA Door Supervisor",
        "SIA CCTV",
        "CCTV monitoring",
        "Access control",
        "Incident reporting",
        "Conflict management",
        "Crowd management",
        "ID verification",
        "Search procedures",
        "Emergency response procedures",
        "Patrolling",
        "Radio communication",
        "Safeguarding",
        "Customer service",
        "Health and safety compliance",
        "GDPR and data protection compliance",
        "Record keeping",
        "Risk assessment",
        "Picking and packing",
        "Order picking",
        "RF scanning",
        "Stock control",
        "Inventory management",
        "Loading and unloading",
        "Forklift driving",
        "Reach truck",
        "Manual handling",
        "PPE compliance",
        "Housekeeping",
        "Room attendant",
        "Front of house",
        "Food hygiene",
        "Bar service",
        "Microsoft Excel",
        "Microsoft Word",
        "Team leadership",
        "Communication",
        "Problem solving",
        "Time management",
    ],
    "variants": {
        "SIA Door Supervisor": ["door supervisor", "sia ds", "ds licence", "ds license", "door sup"],
        "SIA CCTV": ["cctv licence", "cctv license", "sia cctv"],
        "CCTV monitoring": ["cctv", "monitoring cctv", "cctv operator"],
        "Access control": ["access-control", "access", "entry control"],
        "Incident reporting": ["incident report", "incident reports", "incident reporting and record keeping"],
        "GDPR and data protection compliance": ["gdpr", "data protection", "data protection compliance"],
        "Picking and packing": ["picker", "packer", "picking", "packing", "pick and pack"],
        "RF scanning": ["rf scanner", "rf scan", "scan gun", "handheld scanner"],
        "Forklift driving": ["forklift", "forklift driver", "counterbalance"],
        "Reach truck": ["reach", "reach driver", "reach forklift"],
        "Housekeeping": ["house keeper", "housekeeper", "public area cleaner"],
        "Customer service": ["customer support", "client service"],
    },
    "abbreviations": {
        "DS": ["SIA Door Supervisor"],
        "CCTV": ["CCTV monitoring", "SIA CCTV"],
        "GDPR": ["GDPR and data protection compliance"],
        "CRC": ["Criminal record check"],
        "DBS": ["DBS check"],
        "HSE": ["Health and safety compliance"],
    }
}


class SkillCatalogue:
    def __init__(self, data: Dict[str, Any], nlp=None):
        self.data = data
        self.nlp = nlp
        self.canonical: List[str] = data.get("canonical", [])
        self.variants: Dict[str, List[str]] = data.get("variants", {})
        self.abbreviations: Dict[str, List[str]] = data.get("abbreviations", {})

        self.variant_to_canonical: Dict[str, str] = {}
        for canon, vars_ in self.variants.items():
            for v in vars_:
                self.variant_to_canonical[v.lower().strip()] = canon

        self._canon_vectors: Dict[str, np.ndarray] = {}
        if self.nlp and getattr(self.nlp.vocab, "vectors", None) and self.nlp.vocab.vectors.size > 0:
            for canon in self.canonical:
                self._canon_vectors[canon] = self.nlp(canon).vector

    @staticmethod
    def load_from_path(path: Optional[str], nlp=None) -> "SkillCatalogue":
        if path:
            p = Path(path)
            if p.exists() and p.is_file():
                with p.open("r", encoding="utf-8") as f:
                    data = json.load(f)
                return SkillCatalogue(data, nlp=nlp)
        return SkillCatalogue(DEFAULT_SKILL_CATALOGUE, nlp=nlp)

    def canonicalise(self, text: str, context_text: str = "") -> Tuple[str, float, str]:
        t = text.strip()
        if not t:
            return "", 0.0, "passthrough"

        low = t.lower()

        for canon in self.canonical:
            if low == canon.lower():
                return canon, 0.98, "exact"

        if low in self.variant_to_canonical:
            return self.variant_to_canonical[low], 0.95, "variant"

        up = re.sub(r"[^A-Za-z]", "", t).upper()
        if up in self.abbreviations:
            candidates = self.abbreviations[up]
            if len(candidates) == 1:
                return candidates[0], 0.88, "abbrev"
            ctx = context_text.lower()
            if "sia" in ctx or "security" in ctx:
                for c in candidates:
                    if "SIA" in c or "Door Supervisor" in c or "CCTV" in c:
                        return c, 0.80, "abbrev"
            return candidates[0], 0.70, "abbrev"

        if self.nlp and self._canon_vectors:
            v = self.nlp(t).vector
            best = ("", 0.0)
            for canon, cv in self._canon_vectors.items():
                sim = cosine_sim(v, cv)
                if sim > best[1]:
                    best = (canon, sim)
            if best[1] >= CONFIG.embedding_similarity_threshold:
                conf = safe_clip01((best[1] - CONFIG.embedding_similarity_threshold) / (1.0 - CONFIG.embedding_similarity_threshold))
                conf = 0.60 + 0.35 * conf
                return best[0], conf, "embed"

        return self._title_like(t), 0.55, "passthrough"

    @staticmethod
    def _title_like(s: str) -> str:
        s2 = re.sub(r"\s+", " ", s).strip()
        if not s2:
            return s2
        if s2.isupper():
            return " ".join(w.capitalize() for w in s2.split())
        return s2


# ---------------------------------------------------------------------
# Document extraction
# ---------------------------------------------------------------------
class DocumentExtractor:
    def extract_text(self, file_bytes: bytes, filename: str) -> Dict[str, str]:
        name = filename.lower()
        if name.endswith(".pdf"):
            raw = self._extract_pdf(file_bytes)
        elif name.endswith(".docx"):
            raw = self._extract_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {filename}")

        clean = normalise_whitespace_keep_lines(raw)
        return {"raw_text": raw, "clean_text": clean}

    def _extract_docx(self, file_bytes: bytes) -> str:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. pip install python-docx")

        doc = Document(io.BytesIO(file_bytes))
        parts = []

        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)

        for tbl in doc.tables:
            for row in tbl.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    parts.append(row_text)

        return "\n".join(parts).strip()

    def _extract_pdf(self, file_bytes: bytes) -> str:
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber is not installed. pip install pdfplumber")

        pages_text = []
        page_lines_list = []

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                txt = ""
                try:
                    txt = page.extract_text(layout=True, x_tolerance=2, y_tolerance=2) or ""
                except Exception:
                    txt = ""

                if not txt:
                    txt = self._reconstruct_from_words(page) or ""

                pages_text.append(txt)
                lines = [ln.strip() for ln in txt.split("\n") if ln.strip()]
                page_lines_list.append(lines)

        cleaned_pages = self._drop_repeated_header_footer(page_lines_list)
        joined = "\n\n".join("\n".join(lines) for lines in cleaned_pages).strip()

        alpha_chars = sum(ch.isalpha() for ch in joined)
        if CONFIG.enable_ocr_fallback and alpha_chars < CONFIG.min_extracted_alpha_chars_for_pdf:
            ocr_text = self._ocr_pdf(file_bytes)
            if ocr_text and sum(ch.isalpha() for ch in ocr_text) > alpha_chars:
                return ocr_text

        return joined

    @staticmethod
    def _reconstruct_from_words(page) -> str:
        try:
            words = page.extract_words(use_text_flow=True, keep_blank_chars=False)
            if not words:
                return ""
            words_sorted = sorted(words, key=lambda w: (round(w["top"], 1), w["x0"]))
            lines = []
            current_top = None
            current = []
            for w in words_sorted:
                top = round(w["top"], 1)
                if current_top is None:
                    current_top = top
                if abs(top - current_top) > 2.0:
                    lines.append(" ".join(current).strip())
                    current = [w["text"]]
                    current_top = top
                else:
                    current.append(w["text"])
            if current:
                lines.append(" ".join(current).strip())
            return "\n".join([ln for ln in lines if ln.strip()])
        except Exception:
            return ""

    @staticmethod
    def _drop_repeated_header_footer(page_lines_list: List[List[str]]) -> List[List[str]]:
        if not page_lines_list:
            return page_lines_list

        n_pages = len(page_lines_list)
        if n_pages < 2:
            return page_lines_list

        freq: Dict[str, int] = {}
        for lines in page_lines_list:
            unique = set(lines)
            for ln in unique:
                if len(ln) <= 80:
                    freq[ln] = freq.get(ln, 0) + 1

        threshold = max(2, int(math.ceil(n_pages * CONFIG.drop_repeated_header_footer_ratio)))
        to_remove = {ln for ln, c in freq.items() if c >= threshold}

        cleaned = []
        for lines in page_lines_list:
            out = []
            for ln in lines:
                if ln in to_remove:
                    continue
                out.append(ln)
            cleaned.append(out)
        return cleaned

    @staticmethod
    def _ocr_pdf(file_bytes: bytes) -> str:
        if not OCR_AVAILABLE:
            logger.info("OCR fallback requested but OCR dependencies are not installed.")
            return ""

        try:
            images = convert_from_bytes(file_bytes)
            texts = []
            for img in images:
                t = pytesseract.image_to_string(img, lang=CONFIG.ocr_lang)
                if t:
                    texts.append(t)
            return "\n\n".join(texts).strip()
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return ""


# ---------------------------------------------------------------------
# Section parsing
# ---------------------------------------------------------------------
class SectionParser:
    def __init__(self, config: AppConfig):
        self.config = config
        self.header_to_section: Dict[str, str] = {}
        for section, aliases in config.section_aliases.items():
            for a in aliases:
                self.header_to_section[a.lower()] = section

    def split_sections(self, text: str) -> Dict[str, List[str]]:
        lines = [ln.strip() for ln in text.split("\n")]
        sections: Dict[str, List[str]] = {}
        current = "header"
        buf: List[str] = []

        def flush():
            nonlocal buf, current
            if buf:
                sections.setdefault(current, [])
                sections[current].extend([x for x in buf if x.strip()])
            buf = []

        for ln in lines:
            lnl = ln.lower().strip(": ").strip()
            section_found = None

            if looks_like_heading(ln):
                lnl2 = lnl.replace("&", "and")
                for key, sec in self.header_to_section.items():
                    if lnl2 == key:
                        section_found = sec
                        break
                    if lnl2 in key or key in lnl2:
                        if len(lnl2.split()) <= 4:
                            section_found = sec
                            break

            if section_found:
                flush()
                current = section_found
            else:
                buf.append(ln)

        flush()
        return sections


# ---------------------------------------------------------------------
# Contact extraction
# ---------------------------------------------------------------------
class ContactExtractor:
    def __init__(self, config: AppConfig):
        self.config = config

    def extract_emails(self, text: str) -> List[str]:
        found = self.config.email_re.findall(text)
        out = []
        for e in found:
            if EMAIL_AVAILABLE:
                try:
                    v = validate_email(e, check_deliverability=False)
                    out.append(v.email)
                except EmailNotValidError:
                    continue
            else:
                out.append(e)
        return dedupe_preserve_order(out)

    def extract_uk_phones(self, text: str) -> List[str]:
        candidates = []
        raw = re.findall(r"(\+44\s?\d[\d\s]{8,13}|\b0\d[\d\s]{8,13})", text)
        raw = [re.sub(r"\s+", " ", r).strip() for r in raw]
        raw = dedupe_preserve_order(raw)

        if PHONE_AVAILABLE:
            for r in raw:
                try:
                    num = phonenumbers.parse(r, "GB")
                    if phonenumbers.is_possible_number(num) and phonenumbers.is_valid_number(num):
                        fmt = phonenumbers.format_number(num, phonenumbers.PhoneNumberFormat.E164)
                        candidates.append(fmt)
                except Exception:
                    continue
            candidates = dedupe_preserve_order(candidates)
            if candidates:
                return candidates[:3]

        cleaned = []
        for r in raw:
            digits = re.sub(r"\D", "", r)
            if len(digits) >= 10:
                cleaned.append(r)
        return dedupe_preserve_order(cleaned)[:3]

    def extract_postcodes(self, text: str) -> List[str]:
        matches = self.config.uk_postcode_re.findall(text)
        out = []
        for m1, m2 in matches:
            pc = normalise_postcode(m1, m2)
            if pc not in out:
                out.append(pc)
        return out

    def extract_address(self, text: str, postcodes: List[str]) -> Tuple[Optional[str], float, str]:
        if not postcodes:
            return None, 0.0, "none"

        lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
        best = ("", 0.0, "postcode_window")

        for pc in postcodes:
            idxs = [i for i, ln in enumerate(lines) if pc.replace(" ", "") in ln.replace(" ", "").upper()]
            for idx in idxs:
                start = max(0, idx - 3)
                end = min(len(lines), idx + 2)
                chunk = lines[start:end]
                addr = ", ".join([c.strip(" ,") for c in chunk if c.strip()])
                addr = re.sub(r"\s+", " ", addr).strip()

                score = 0.4
                if re.search(r"\b(road|rd|street|st|lane|ln|avenue|ave|drive|dr|close|cl|way|court|ct|flat|house)\b", addr, re.I):
                    score += 0.25
                if re.search(r"\b\d{1,4}\b", addr):
                    score += 0.15
                if len(addr) >= 20:
                    score += 0.10
                score = safe_clip01(score)

                if score > best[1]:
                    best = (addr, score, "postcode_window")

        if best[1] <= 0.0:
            return None, 0.0, "none"
        return best


# ---------------------------------------------------------------------
# Name extraction and reconciliation
# ---------------------------------------------------------------------
class NameResolver:
    def __init__(self, nlp):
        self.nlp = nlp

    @staticmethod
    def infer_name_from_filename(filename: str) -> Tuple[Optional[str], float]:
        stem = Path(filename).stem
        stem = re.sub(r"(?i)\b(resume|cv)\b", " ", stem)
        stem = re.sub(r"(?i)\b(final|latest|updated)\b", " ", stem)
        stem = re.sub(r"[_\.\-]+", " ", stem)
        stem = re.sub(r"([a-z])([A-Z])", r"\1 \2", stem)
        stem = re.sub(r"\b\d+\b", " ", stem)
        stem = re.sub(r"\s+", " ", stem).strip()

        if not stem:
            return None, 0.0

        tokens = stem.split()
        if len(tokens) < 2:
            return None, 0.0
        if len(tokens) > 6:
            tokens = tokens[:6]

        name = " ".join([t.capitalize() if not t.isupper() else t for t in tokens]).strip()
        if re.search(r"[^A-Za-z\s'’\-]", name):
            return None, 0.0

        conf = 0.65
        if 2 <= len(tokens) <= 4:
            conf = 0.80
        return name, conf

    def extract_name_from_content(self, text: str) -> Tuple[Optional[str], float]:
        if not text.strip():
            return None, 0.0

        lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
        header_lines = lines[:40]
        header_block = "\n".join(header_lines)

        bad_keywords = {
            "curriculum vitae", "cv", "resume", "profile", "summary",
            "contact", "details", "email", "e-mail", "phone", "mobile", "tel",
            "address", "postcode", "linkedin", "github", "nationality", "dob", "date of birth"
        }

        def strip_noise(s: str) -> str:
            s = re.sub(CONFIG.email_re, " ", s)
            s = re.sub(r"(\+44\s?\d[\d\s]{8,13}|\b0\d[\d\s]{8,13})", " ", s)
            s = re.sub(r"(?:https?://)?(?:www\.)?linkedin\.com/\S+", " ", s, flags=re.I)
            s = re.sub(r"(?:https?://)?(?:www\.)?github\.com/\S+", " ", s, flags=re.I)
            s = re.sub(r"\s+", " ", s).strip()
            return s

        def plausible_name(s: str) -> bool:
            if not s:
                return False
            if any(k in s.lower() for k in bad_keywords):
                return False
            if len(s) > 60:
                return False
            if re.search(r"[^A-Za-z\s'’\-\.]", s):
                return False
            parts = s.replace(".", "").split()
            if len(parts) < 2 or len(parts) > 5:
                return False
            if any(len(p) == 1 for p in parts):
                if sum(len(p) == 1 for p in parts) >= 2:
                    return False
            return True

        candidates: List[Tuple[str, float]] = []
        name_line_re = re.compile(r"^([A-Za-z][A-Za-z'’\-]+(?:\s+[A-Za-z][A-Za-z'’\-]+){1,4})$")

        for i, ln in enumerate(header_lines[:15]):
            ln2 = strip_noise(ln)
            if not ln2:
                continue
            if any(k in ln2.lower() for k in bad_keywords):
                continue
            m = name_line_re.match(ln2)
            if m:
                nm = m.group(1).strip()
                score = 0.78 + max(0.0, (10 - i) * 0.01)
                candidates.append((nm, safe_clip01(score)))

            m2 = re.search(r"(?i)\bname\s*[:\-]\s*([A-Za-z][A-Za-z'’\-]+(?:\s+[A-Za-z][A-Za-z'’\-]+){1,4})\b", ln)
            if m2:
                nm = m2.group(1).strip()
                candidates.append((nm, 0.88))

        if self.nlp:
            try:
                doc = self.nlp(header_block)
                for ent in doc.ents:
                    if ent.label_ == "PERSON":
                        nm = strip_noise(ent.text).strip()
                        if nm.isupper():
                            nm = " ".join(w.capitalize() for w in nm.split())
                        if plausible_name(nm):
                            candidates.append((nm, 0.90))
            except Exception:
                pass

        if not candidates:
            return None, 0.0

        candidates.sort(key=lambda x: (x[1], len(x[0].split())), reverse=True)
        best = candidates[0]
        return best[0], best[1]

    def reconcile(self, content_name: Optional[str], content_conf: float, filename: str) -> Dict[str, Any]:
        file_name, file_conf = self.infer_name_from_filename(filename)

        if content_name and not file_name:
            return {"full_name": content_name, "name_source": "content", "name_confidence": safe_clip01(content_conf)}
        if file_name and not content_name:
            return {"full_name": file_name, "name_source": "filename", "name_confidence": safe_clip01(file_conf)}
        if not content_name and not file_name:
            return {"full_name": None, "name_source": "none", "name_confidence": 0.0}

        c = content_name or ""
        f = file_name or ""

        if RAPIDFUZZ_AVAILABLE:
            sim = fuzz.token_set_ratio(c, f) / 100.0
        else:
            sim = self._simple_token_similarity(c, f)

        if sim >= 0.90:
            conf = safe_clip01(max(content_conf, file_conf) + 0.05)
            return {"full_name": content_name, "name_source": "merged", "name_confidence": conf}
        if sim >= 0.75:
            merged = self._merge_names(c, f)
            conf = safe_clip01(max(content_conf, file_conf) - 0.05)
            return {"full_name": merged, "name_source": "merged", "name_confidence": conf}

        if content_conf >= file_conf:
            return {"full_name": content_name, "name_source": "content", "name_confidence": safe_clip01(content_conf - 0.10)}
        return {"full_name": file_name, "name_source": "filename", "name_confidence": safe_clip01(file_conf - 0.10)}

    @staticmethod
    def _simple_token_similarity(a: str, b: str) -> float:
        ta = set(a.lower().split())
        tb = set(b.lower().split())
        if not ta or not tb:
            return 0.0
        return len(ta & tb) / len(ta | tb)

    @staticmethod
    def _merge_names(a: str, b: str) -> str:
        ta = a.split()
        tb = b.split()
        out = []
        seen = set()
        for t in ta + tb:
            k = t.lower()
            if k not in seen:
                seen.add(k)
                out.append(t)
        return " ".join(out)


# ---------------------------------------------------------------------
# Licences and certifications
# ---------------------------------------------------------------------
class LicenceExtractor:
    def __init__(self, catalogue: SkillCatalogue):
        self.catalogue = catalogue
        self.keywords = [
            "sia", "door supervisor", "cctv", "first aid", "emergency first aid",
            "crc", "dbs", "cscs", "haccp", "food hygiene", "licence", "license", "certificate", "certification"
        ]

    def extract(self, sections: Dict[str, List[str]], full_text: str) -> List[Dict[str, Any]]:
        results: List[Dict[str, Any]] = []
        source_section = "certifications" if "certifications" in sections else None
        lines = sections.get("certifications", [])
        if not lines:
            lines = [ln.strip() for ln in full_text.split("\n") if ln.strip()]
            source_section = "content"

        for ln in lines:
            low = ln.lower()
            if len(ln) < 3:
                continue
            if any(k in low for k in self.keywords):
                canon, conf, method = self.catalogue.canonicalise(ln, context_text=full_text)
                results.append({
                    "text": ln,
                    "canonical": canon,
                    "confidence": safe_clip01(conf),
                    "source": source_section,
                    "method": method,
                })

        seen = set()
        out = []
        for r in results:
            key = (r["canonical"] or r["text"]).lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(r)
        return out


# ---------------------------------------------------------------------
# Skills extraction (kept as in your version)
# ---------------------------------------------------------------------
class SkillsExtractor:
    def __init__(self, nlp, catalogue: SkillCatalogue):
        self.nlp = nlp
        self.catalogue = catalogue
        self.stop_starts = {"responsible", "responsibilities", "duties", "role", "worked", "working"}

    def extract(self, sections: Dict[str, List[str]], full_text: str) -> Dict[str, Any]:
        mentions: List[Dict[str, Any]] = []

        skill_lines = sections.get("skills", [])
        mentions.extend(self._extract_from_lines(skill_lines, section="skills", context_text=full_text, base_conf=0.90))

        summary_lines = sections.get("summary", [])
        mentions.extend(self._extract_from_lines(summary_lines, section="summary", context_text=full_text, base_conf=0.70))

        exp_lines = sections.get("experience", [])
        exp_candidates = [ln for ln in exp_lines if CONFIG.bullet_re.match(ln) or (len(ln.split()) <= 12 and len(ln) >= 8)]
        mentions.extend(self._extract_from_lines(exp_candidates, section="experience", context_text=full_text, base_conf=0.65))

        mentions.extend(self._extract_dictionary_hits(full_text))

        canon_all = []
        for m in mentions:
            if m.get("canonical"):
                canon_all.append(m["canonical"])
        canon_all = dedupe_preserve_order(canon_all)

        if CONFIG.keep_unmatched_skill_phrases:
            for m in mentions:
                if m.get("canonical"):
                    continue
                if m.get("section") == "skills":
                    canon_all.append(m["text"])
            canon_all = dedupe_preserve_order(canon_all)

        seen = set()
        dedup_mentions = []
        for m in mentions:
            key = (m.get("canonical", "").lower(), m.get("text", "").lower(), m.get("section", "").lower())
            if key in seen:
                continue
            seen.add(key)
            dedup_mentions.append(m)

        dedup_mentions.sort(key=lambda x: x.get("confidence", 0.0), reverse=True)

        return {
            "all_skills": canon_all,
            "mentions": dedup_mentions
        }

    def _extract_from_lines(self, lines: List[str], section: str, context_text: str, base_conf: float) -> List[Dict[str, Any]]:
        out: List[Dict[str, Any]] = []
        for ln in lines:
            ln0 = ln.strip()
            if not ln0:
                continue
            ln0 = CONFIG.bullet_re.sub("", ln0).strip()

            candidates = self._split_skill_line(ln0)
            candidates.extend(self._phrase_mine(ln0))
            candidates = dedupe_preserve_order([c for c in candidates if len(c) >= 2])

            for c in candidates:
                canon, conf, method = self.catalogue.canonicalise(c, context_text=context_text)
                conf2 = conf
                if section == "skills":
                    conf2 = max(conf2, base_conf if method in {"exact", "variant", "abbrev"} else base_conf - 0.10)
                else:
                    conf2 = max(conf2, base_conf - 0.10)

                out.append({
                    "text": c,
                    "canonical": canon if method != "passthrough" else (canon if not self._too_generic(canon) else ""),
                    "section": section,
                    "sentence": ln.strip(),
                    "confidence": safe_clip01(conf2),
                    "method": method
                })
        return out

    def _extract_dictionary_hits(self, full_text: str) -> List[Dict[str, Any]]:
        out: List[Dict[str, Any]] = []
        low = full_text.lower()

        for canon in self.catalogue.canonical:
            if canon.lower() in low:
                out.append({
                    "text": canon,
                    "canonical": canon,
                    "section": "content",
                    "sentence": canon,
                    "confidence": 0.85,
                    "method": "dictionary"
                })

        for v, canon in self.catalogue.variant_to_canonical.items():
            if v in low:
                out.append({
                    "text": v,
                    "canonical": canon,
                    "section": "content",
                    "sentence": v,
                    "confidence": 0.82,
                    "method": "dictionary"
                })

        for abbr, canons in self.catalogue.abbreviations.items():
            if re.search(rf"\b{re.escape(abbr)}\b", full_text, flags=re.I):
                for c in canons[:1]:
                    out.append({
                        "text": abbr,
                        "canonical": c,
                        "section": "content",
                        "sentence": abbr,
                        "confidence": 0.75,
                        "method": "abbrev"
                    })

        return out

    @staticmethod
    def _split_skill_line(line: str) -> List[str]:
        parts = [line.strip()]
        chunks = re.split(r"[;,/|]+", line)
        for ch in chunks:
            ch2 = ch.strip()
            if ch2 and ch2.lower() not in {"skills", "technical skills"}:
                parts.append(ch2)

        chunks2 = [c.strip() for c in re.split(r"\s*,\s*", line) if c.strip()]
        parts.extend(chunks2)

        if " and " in line.lower():
            and_parts = [p.strip() for p in re.split(r"(?i)\band\b", line) if p.strip()]
            if 2 <= len(and_parts) <= 4:
                parts.extend(and_parts)

        return [p for p in parts if len(p.split()) <= 14]

    def _phrase_mine(self, line: str) -> List[str]:
        out = []
        if not self.nlp:
            return out

        try:
            doc = self.nlp(line)
            for nc in getattr(doc, "noun_chunks", []):
                t = nc.text.strip().strip(" .,:;()[]{}")
                if 2 <= len(t.split()) <= 12:
                    out.append(t)

            buff = []
            for tok in doc:
                if tok.is_stop and tok.text.lower() not in {"and", "&"}:
                    if len(buff) >= 2:
                        out.append(" ".join(buff))
                    buff = []
                    continue

                if tok.pos_ in {"NOUN", "PROPN", "ADJ"} or tok.text in {"&", "and"}:
                    buff.append(tok.text)
                else:
                    if len(buff) >= 2:
                        out.append(" ".join(buff))
                    buff = []
            if len(buff) >= 2:
                out.append(" ".join(buff))

        except Exception:
            return out

        cleaned = []
        for t in out:
            t2 = re.sub(r"\s+", " ", t).strip()
            if not t2:
                continue
            if t2.lower().split()[0] in self.stop_starts:
                continue
            if self._too_generic(t2):
                continue
            cleaned.append(t2)
        return cleaned

    @staticmethod
    def _too_generic(s: str) -> bool:
        low = s.lower().strip()
        if low in {"skills", "experience", "responsibilities", "duties"}:
            return True
        if len(low.split()) == 1 and low in {"work", "team", "company", "people"}:
            return True
        return False


# ---------------------------------------------------------------------
# Experience extraction (kept as in your version)
# ---------------------------------------------------------------------
class ExperienceExtractor:
    def __init__(self, nlp, config: AppConfig):
        self.nlp = nlp
        self.config = config

    def extract(self, sections: Dict[str, List[str]], full_text: str) -> List[Dict[str, Any]]:
        exp_lines = sections.get("experience", [])
        if not exp_lines:
            exp_lines = [ln.strip() for ln in full_text.split("\n") if ln.strip()]

        blocks = self._split_into_role_blocks(exp_lines)
        roles = []
        for b in blocks:
            role = self._parse_role_block(b)
            if role:
                roles.append(role)
        return roles

    def _split_into_role_blocks(self, lines: List[str]) -> List[List[str]]:
        blocks: List[List[str]] = []
        buf: List[str] = []

        def flush():
            nonlocal buf
            if buf:
                cleaned = [x for x in buf if x.strip()]
                if cleaned:
                    blocks.append(cleaned)
            buf = []

        for ln in lines:
            if not ln.strip():
                if len(buf) >= 3:
                    flush()
                else:
                    buf.append(ln)
                continue

            if self.config.date_range_re.search(ln) and len(buf) >= 2:
                flush()
                buf.append(ln)
            else:
                buf.append(ln)

        flush()

        merged: List[List[str]] = []
        for b in blocks:
            if len(b) <= 2 and merged:
                merged[-1].extend(b)
            else:
                merged.append(b)
        return merged

    def _parse_role_block(self, lines: List[str]) -> Optional[Dict[str, Any]]:
        raw_block = [ln.strip() for ln in lines if ln.strip()]
        if len(raw_block) < 2:
            return None

        date_line_idx = None
        date_match = None
        for i, ln in enumerate(raw_block[:8]):
            m = self.config.date_range_re.search(ln)
            if m:
                date_line_idx = i
                date_match = m
                break

        header_lines = raw_block[:date_line_idx] if date_line_idx is not None else raw_block[:2]
        header_text = " | ".join(header_lines).strip()

        start_resp_idx = (date_line_idx + 1) if date_line_idx is not None else len(header_lines)
        resp_lines = raw_block[start_resp_idx:]
        responsibilities = self._clean_bullets(resp_lines)

        start_date, end_date, date_conf = self._parse_date_range(date_match.group(0)) if date_match else (None, None, 0.0)
        job_title, employer, location, header_conf, evidence = self._extract_header_fields(header_text)
        industry_tags = self._infer_industry(job_title, employer, responsibilities)
        dur = self._duration_months(start_date, end_date)

        role = {
            "job_title": job_title,
            "employer": employer,
            "location": location,
            "start_date": start_date,
            "end_date": end_date,
            "duration_months": dur,
            "responsibilities": responsibilities,
            "industry_tags": industry_tags,
            "evidence": evidence,
            "confidence": {
                "job_title": safe_clip01(header_conf.get("job_title", 0.0)),
                "employer": safe_clip01(header_conf.get("employer", 0.0)),
                "location": safe_clip01(header_conf.get("location", 0.0)),
                "dates": safe_clip01(date_conf),
                "responsibilities": safe_clip01(0.70 if responsibilities else 0.30),
                "overall": safe_clip01(
                    0.25 * header_conf.get("job_title", 0.0) +
                    0.25 * header_conf.get("employer", 0.0) +
                    0.20 * date_conf +
                    0.10 * (0.70 if responsibilities else 0.30) +
                    0.20
                )
            }
        }

        if not role["job_title"] and not role["employer"] and not date_match:
            return None
        return role

    @staticmethod
    def _clean_bullets(lines: List[str]) -> List[str]:
        out = []
        for ln in lines:
            t = ln.strip()
            if not t:
                continue
            t = CONFIG.bullet_re.sub("", t).strip()
            if looks_like_heading(t):
                continue
            if len(t) < 8:
                continue
            out.append(t)
        return dedupe_preserve_order(out)

    def _parse_date_range(self, s: str) -> Tuple[Optional[str], Optional[str], float]:
        m = self.config.date_range_re.search(s)
        if not m:
            return None, None, 0.0

        start_raw = m.group("start")
        end_raw = m.group("end")

        sd, sc = parse_month_year(start_raw)
        if sd is None:
            return None, None, 0.0

        end_lower = end_raw.strip().lower()
        if end_lower in {"present", "current", "now"}:
            return date_to_yyyy_mm(sd), "Current", 0.90

        ed, ec = parse_month_year(end_raw)
        if ed is None:
            if re.fullmatch(r"(19|20)\d{2}", end_raw.strip()):
                try:
                    y = int(end_raw.strip())
                    ed = date(y, 1, 1)
                    ec = 0.70
                except Exception:
                    ed = None

        if ed is None:
            return date_to_yyyy_mm(sd), None, 0.70

        conf = min(0.92, 0.60 + 0.20 * sc + 0.20 * ec)
        return date_to_yyyy_mm(sd), date_to_yyyy_mm(ed), conf

    def _extract_header_fields(self, header_text: str) -> Tuple[Optional[str], Optional[str], Optional[str], Dict[str, float], List[str]]:
        evidence = [header_text]
        conf = {"job_title": 0.55, "employer": 0.45, "location": 0.35}

        parts = [p.strip() for p in re.split(r"\s*\|\s*|\s*-\s*|\s*•\s*|\s*,\s*", header_text) if p.strip()]
        parts = parts[:6]

        job_title = None
        employer = None
        location = None

        if self.nlp:
            try:
                doc = self.nlp(header_text)
                orgs = [ent.text.strip() for ent in doc.ents if ent.label_ in {"ORG"}]
                geos = [ent.text.strip() for ent in doc.ents if ent.label_ in {"GPE", "LOC"}]
                if orgs:
                    employer = orgs[0]
                    conf["employer"] = 0.75
                if geos:
                    location = geos[0]
                    conf["location"] = 0.70
            except Exception:
                pass

        for p in parts:
            pl = p.lower()
            if any(k in pl for k in ["road", "street", "lane", "avenue", "drive"]) and re.search(r"\d", p):
                continue
            if len(p.split()) <= 8:
                if any(w in pl for w in ["supervisor", "officer", "guard", "steward", "operative", "picker", "packer", "driver", "housekeeper", "waiter", "reception", "admin", "security"]):
                    job_title = p
                    conf["job_title"] = 0.80
                    break

        if not job_title and parts:
            job_title = parts[0]
            conf["job_title"] = 0.65

        if not employer:
            for p in parts[1:]:
                if len(p.split()) <= 10 and any(ch.isalpha() for ch in p):
                    if re.search(r"\b(ltd|limited|group|services|security|solutions|hotel|warehouse)\b", p, re.I):
                        employer = p
                        conf["employer"] = 0.70
                        break
            if not employer and len(parts) >= 2:
                employer = parts[1]
                conf["employer"] = 0.55

        if not location:
            for p in parts:
                if re.search(r"\b(uk|london|manchester|birmingham|leeds|glasgow|bristol|hounslow|slough|reading)\b", p, re.I):
                    location = p
                    conf["location"] = 0.55
                    break

        return job_title, employer, location, conf, evidence

    def _infer_industry(self, job_title: Optional[str], employer: Optional[str], responsibilities: List[str]) -> List[str]:
        text = " ".join([job_title or "", employer or ""] + responsibilities).lower()
        tags = []
        for k, kws in self.config.industry_keywords.items():
            if any(kw in text for kw in kws):
                tags.append(k)
        return dedupe_preserve_order(tags)

    @staticmethod
    def _duration_months(start_date: Optional[str], end_date: Optional[str]) -> Optional[int]:
        if not start_date:
            return None
        try:
            sy, sm = [int(x) for x in start_date.split("-")]
            sd = date(sy, sm, 1)
        except Exception:
            return None

        if not end_date or end_date == "Current":
            return None

        try:
            ey, em = [int(x) for x in end_date.split("-")]
            ed = date(ey, em, 1)
        except Exception:
            return None

        if ed < sd:
            return None
        return months_between(sd, ed) + 1


# ---------------------------------------------------------------------
# Keywords and summary
# ---------------------------------------------------------------------
class KeywordSummariser:
    def __init__(self):
        self.stop = {
            "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for",
            "of", "with", "by", "from", "as", "is", "was", "are", "were", "been",
            "this", "that", "these", "those", "i", "we", "they", "he", "she", "you",
        }

    def keywords(self, text: str, top_n: int = 35) -> List[str]:
        low = text.lower()
        words = re.findall(r"\b[a-z]{3,}\b", low)
        words = [w for w in words if w not in self.stop]
        if not words:
            return []
        freq = {}
        for w in words:
            freq[w] = freq.get(w, 0) + 1
        items = sorted(freq.items(), key=lambda x: x[1], reverse=True)
        return [k for k, _ in items[:top_n]]

    def summary(self, text: str, keywords: List[str], max_words: int = 160) -> str:
        lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
        scored = []
        top_kw = set([k.lower() for k in keywords[:18]])
        for ln in lines[:120]:
            if len(ln.split()) < 6:
                continue
            score = sum(1 for k in top_kw if k in ln.lower())
            if score > 0:
                scored.append((ln, score))
        scored.sort(key=lambda x: x[1], reverse=True)

        out = []
        wc = 0
        for ln, _ in scored[:6]:
            n = len(ln.split())
            if wc + n > max_words:
                continue
            out.append(ln)
            wc += n
        if out:
            return " ".join(out)
        return " ".join(text.split()[:max_words])


# ---------------------------------------------------------------------
# Main CV parser
# Adds parse_bytes for WorkDrive downloads
# ---------------------------------------------------------------------
class CVParser:
    def __init__(self, nlp, catalogue: SkillCatalogue, config: AppConfig):
        self.config = config
        self.docx = DocumentExtractor()
        self.sections = SectionParser(config)
        self.contact = ContactExtractor(config)
        self.name = NameResolver(nlp)
        self.skills = SkillsExtractor(nlp, catalogue)
        self.licences = LicenceExtractor(catalogue)
        self.exp = ExperienceExtractor(nlp, config)
        self.kw = KeywordSummariser()

    def parse_path(self, path: Path) -> Dict[str, Any]:
        file_bytes = path.read_bytes()
        return self.parse_bytes(file_bytes=file_bytes, filename=path.name, path_hint=str(path))

    def parse_bytes(self, file_bytes: bytes, filename: str, path_hint: str = "") -> Dict[str, Any]:
        t0 = time.time()
        extracted = self.docx.extract_text(file_bytes, filename)
        clean_text = extracted["clean_text"]

        sections = self.sections.split_sections(clean_text)

        content_name, content_conf = self.name.extract_name_from_content(clean_text)
        name_info = self.name.reconcile(content_name, content_conf, filename)

        emails = self.contact.extract_emails(clean_text)
        phones = self.contact.extract_uk_phones(clean_text)
        postcodes = self.contact.extract_postcodes(clean_text)
        addr_val, addr_conf, addr_src = self.contact.extract_address(clean_text, postcodes)

        skills = self.skills.extract(sections, clean_text)
        licences = self.licences.extract(sections, clean_text)
        experiences = self.exp.extract(sections, clean_text)

        keywords = self.kw.keywords(clean_text, top_n=35)
        summary = self.kw.summary(clean_text, keywords, max_words=160)

        elapsed = time.time() - t0

        result = {
            "metadata": {
                "filename": filename,
                "path": path_hint,
                "processed_at": datetime.now().isoformat(),
                "elapsed_seconds": round(elapsed, 3),
                "text_length": len(clean_text),
            },
            "personal_information": {
                "full_name": name_info["full_name"],
                "name_source": name_info["name_source"],
                "name_confidence": name_info["name_confidence"],
            },
            "contact_information": {
                "email": {
                    "value": emails[0] if emails else None,
                    "confidence": 0.95 if emails else 0.0,
                    "source": "content"
                },
                "all_emails": emails,
                "phone": {
                    "value": phones[0] if phones else None,
                    "confidence": 0.90 if phones else 0.0,
                    "source": "content"
                },
                "all_phones": phones,
                "postcode": {
                    "value": postcodes[0] if postcodes else None,
                    "confidence": 0.92 if postcodes else 0.0,
                    "source": "content"
                },
                "all_postcodes": postcodes,
                "address": {
                    "value": addr_val,
                    "confidence": addr_conf,
                    "source": addr_src
                },
            },
            "skills": skills,
            "professional_experience": experiences,
            "licences_and_certifications": licences,
            "keywords": keywords,
            "summary": summary,
            "insights": {
                "total_experience_entries": len(experiences),
                "total_skills": len(skills.get("all_skills", [])),
                "total_licences_and_certifications": len(licences),
                "has_email": bool(emails),
                "has_phone": bool(phones),
                "has_postcode": bool(postcodes),
                "has_address": bool(addr_val),
                "completeness_score": self._completeness_score(name_info, emails, phones, postcodes, addr_val, experiences, skills, licences)
            }
        }
        return result

    @staticmethod
    def _completeness_score(name_info, emails, phones, postcodes, addr_val, experiences, skills, licences) -> int:
        score = 0
        if name_info.get("full_name"):
            score += 15
        if emails:
            score += 15
        if phones:
            score += 15
        if postcodes:
            score += 10
        if addr_val:
            score += 10
        if experiences:
            score += 20
        if skills.get("all_skills") and len(skills["all_skills"]) >= 8:
            score += 10
        if licences:
            score += 5
        return min(100, score)


# ---------------------------------------------------------------------
# Local folder scanning
# Searches every subfolder recursively and only returns Resume/CV prefix files
# ---------------------------------------------------------------------
def discover_cv_files_local(root_dir: str, config: AppConfig, exact_filename: str = "") -> List[Path]:
    root = Path(root_dir).expanduser().resolve()
    if not root.exists() or not root.is_dir():
        raise ValueError("Folder path does not exist or is not a directory.")

    out: List[Path] = []
    max_bytes = config.max_file_mb * 1024 * 1024

    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if not d.startswith(".")]

        for fn in filenames:
            if not is_resume_filename(fn, supported_exts=config.supported_exts, exact_filename=exact_filename):
                continue

            p = Path(dirpath) / fn
            try:
                if p.stat().st_size > max_bytes:
                    continue
            except Exception:
                continue

            out.append(p)

    out.sort(key=lambda x: str(x).lower())
    return out


# ---------------------------------------------------------------------
# Streamlit UI
# ---------------------------------------------------------------------
def main():
    st.set_page_config(page_title="CV Parser WorkDrive Recursive", page_icon="📄", layout="wide")

    st.title("📄 CV Parser WorkDrive Recursive")
    st.caption("Recursively scans folder, subfolder, subfolder of subfolder, then filters only Resume or CV filename prefix.")

    nlp = load_spacy_model()
    if not nlp:
        st.warning("spaCy model not loaded. Parser will run but accuracy will be lower.")

    with st.sidebar:
        st.header("⚙️ Source")
        source = st.radio("Choose source", ["Zoho WorkDrive", "Local folder"], index=0)

        st.markdown("---")
        st.header("🔎 File filtering")
        exact_filename = st.text_input(
            "Exact filename (optional)",
            value="",
            help="If you enter: Resume-Manisha_mukeshbhai-Radadiya.docx then only this file will be returned if found anywhere in subfolders."
        )
        include_subfolders = st.checkbox("Search all subfolders", value=True)

        st.markdown("---")
        st.header("🧠 Parser settings")
        enable_ocr = st.checkbox("Enable OCR fallback for scanned PDFs", value=CONFIG.enable_ocr_fallback)
        CONFIG.enable_ocr_fallback = enable_ocr

        sim_threshold = st.slider(
            "Embedding similarity threshold",
            min_value=0.55,
            max_value=0.90,
            value=float(CONFIG.embedding_similarity_threshold),
            step=0.01
        )
        CONFIG.embedding_similarity_threshold = float(sim_threshold)

        show_debug = st.checkbox("Show debug logs", value=False)

        st.markdown("---")
        st.write("Resume filter rule")
        st.write("Extensions:", ", ".join(CONFIG.supported_exts))
        st.write("Filename must start with Resume or CV")
        st.write("Example:", "Resume-Manisha_mukeshbhai-Radadiya.docx")

    # Source specific inputs
    workdrive_folder_link_or_id = ""
    local_folder = ""

    if source == "Zoho WorkDrive":
        st.subheader("Zoho WorkDrive")
        workdrive_folder_link_or_id = st.text_input(
            "WorkDrive folder link or folder ID",
            value="",
            help="Paste the folder link or the folder ID. This app will extract the ID if possible."
        )
    else:
        st.subheader("Local folder")
        local_folder = st.text_input("Local folder path", value=str(Path.cwd()))

    catalogue_path = st.text_input("Skills catalogue JSON path (optional)", value="")
    catalogue = SkillCatalogue.load_from_path(catalogue_path.strip() or None, nlp=nlp)
    parser = CVParser(nlp, catalogue, CONFIG)

    col1, col2 = st.columns([1, 1])

    # Scan button
    with col1:
        if st.button("🔎 Scan and list resume files", type="secondary"):
            try:
                if source == "Local folder":
                    files = discover_cv_files_local(local_folder, CONFIG, exact_filename=exact_filename.strip())
                    st.session_state["matched_local_files"] = files
                    st.session_state["matched_workdrive_files"] = []
                    st.success(f"Found {len(files)} resume files (recursive).")
                else:
                    if not workdrive_folder_link_or_id.strip():
                        st.error("Paste a WorkDrive folder link or folder ID.")
                    else:
                        wd = ZohoWorkDriveClient()
                        folder_id = wd.extract_resource_id(workdrive_folder_link_or_id)
                        items = wd.list_resume_files_recursive(
                            root_folder_id=folder_id,
                            include_subfolders=include_subfolders,
                            supported_exts=CONFIG.supported_exts,
                            exact_filename=exact_filename.strip()
                        )
                        st.session_state["matched_workdrive_files"] = items
                        st.session_state["matched_local_files"] = []
                        st.success(f"Found {len(items)} resume files (recursive).")
            except Exception as e:
                st.error(str(e))
                if show_debug:
                    st.code(traceback.format_exc())

    # Process button
    with col2:
        if st.button("🚀 Process matched files", type="primary"):
            try:
                results = []
                progress = st.progress(0)
                status = st.empty()

                if source == "Local folder":
                    files: List[Path] = discover_cv_files_local(local_folder, CONFIG, exact_filename=exact_filename.strip())
                    if not files:
                        st.warning("No matching resume files found.")
                    else:
                        for i, f in enumerate(files, start=1):
                            status.write(f"Processing {i}/{len(files)}: {f.name}")
                            try:
                                results.append(parser.parse_path(f))
                            except Exception as ex:
                                results.append({
                                    "metadata": {"filename": f.name, "path": str(f), "processed_at": datetime.now().isoformat()},
                                    "error": str(ex)
                                })
                            progress.progress(int(i * 100 / len(files)))

                else:
                    if not workdrive_folder_link_or_id.strip():
                        st.error("Paste a WorkDrive folder link or folder ID.")
                    else:
                        wd = ZohoWorkDriveClient()
                        folder_id = wd.extract_resource_id(workdrive_folder_link_or_id)
                        items = wd.list_resume_files_recursive(
                            root_folder_id=folder_id,
                            include_subfolders=include_subfolders,
                            supported_exts=CONFIG.supported_exts,
                            exact_filename=exact_filename.strip()
                        )
                        if not items:
                            st.warning("No matching resume files found in WorkDrive.")
                        else:
                            for i, it in enumerate(items, start=1):
                                status.write(f"Downloading and processing {i}/{len(items)}: {it['virtual_path']}")
                                try:
                                    b = wd.download_file_bytes(it["id"])
                                    results.append(parser.parse_bytes(
                                        file_bytes=b,
                                        filename=it["name"],
                                        path_hint=f"workdrive:{it['virtual_path']}"
                                    ))
                                except Exception as ex:
                                    results.append({
                                        "metadata": {"filename": it.get("name"), "path": f"workdrive:{it.get('virtual_path')}", "processed_at": datetime.now().isoformat()},
                                        "error": str(ex)
                                    })
                                progress.progress(int(i * 100 / len(items)))

                st.session_state["parsed_results"] = results
                st.success("Processing complete.")
            except Exception as e:
                st.error(str(e))
                if show_debug:
                    st.code(traceback.format_exc())

    st.markdown("---")

    # Show matched files
    if source == "Local folder":
        matched = st.session_state.get("matched_local_files", [])
        st.subheader("Matched resume files (local)")
        if matched:
            st.write(f"Total: {len(matched)}")
            st.dataframe(
                pd.DataFrame({"file": [f.name for f in matched], "path": [str(f) for f in matched]}),
                use_container_width=True
            )
        else:
            st.info("No matched files yet. Click Scan.")
    else:
        matched = st.session_state.get("matched_workdrive_files", [])
        st.subheader("Matched resume files (WorkDrive)")
        if matched:
            st.write(f"Total: {len(matched)}")
            st.dataframe(
                pd.DataFrame({
                    "file": [x.get("name") for x in matched],
                    "path": [x.get("virtual_path") for x in matched],
                    "id": [x.get("id") for x in matched],
                }),
                use_container_width=True
            )
        else:
            st.info("No matched files yet. Click Scan.")

    st.markdown("---")

    # Results
    results = st.session_state.get("parsed_results", [])
    if results:
        st.subheader("Results overview")

        rows = []
        for r in results:
            if "error" in r:
                rows.append({
                    "filename": r.get("metadata", {}).get("filename"),
                    "name": None,
                    "email": None,
                    "phone": None,
                    "postcode": None,
                    "skills_count": 0,
                    "exp_count": 0,
                    "completeness": 0,
                    "error": r.get("error")
                })
                continue

            rows.append({
                "filename": r["metadata"]["filename"],
                "name": r["personal_information"]["full_name"],
                "email": r["contact_information"]["email"]["value"],
                "phone": r["contact_information"]["phone"]["value"],
                "postcode": r["contact_information"]["postcode"]["value"],
                "skills_count": len(r.get("skills", {}).get("all_skills", [])),
                "exp_count": len(r.get("professional_experience", [])),
                "completeness": r.get("insights", {}).get("completeness_score", 0),
                "error": ""
            })

        df = pd.DataFrame(rows)
        st.dataframe(df, use_container_width=True)

        st.markdown("---")
        st.subheader("Detailed outputs")

        for idx, r in enumerate(results, start=1):
            fname = r.get("metadata", {}).get("filename", f"cv_{idx}")
            with st.expander(f"{idx}. {fname}", expanded=(idx == 1)):
                if "error" in r:
                    st.error(r["error"])
                    st.code(json.dumps(r, indent=2), language="json")
                else:
                    c1, c2, c3 = st.columns(3)
                    with c1:
                        st.write("**Name**", r["personal_information"]["full_name"])
                        st.write("Source:", r["personal_information"]["name_source"])
                        st.write("Confidence:", r["personal_information"]["name_confidence"])
                    with c2:
                        st.write("**Email**", r["contact_information"]["email"]["value"])
                        st.write("**Phone**", r["contact_information"]["phone"]["value"])
                    with c3:
                        st.write("**Postcode**", r["contact_information"]["postcode"]["value"])
                        st.write("**Completeness**", r["insights"]["completeness_score"])
                        st.progress(r["insights"]["completeness_score"] / 100)

                    st.write("**Top skills**")
                    st.write(", ".join(r.get("skills", {}).get("all_skills", [])[:25]))

                    st.write("**Experience**")
                    exps = r.get("professional_experience", [])
                    for e in exps[:5]:
                        st.write(
                            f"- {e.get('job_title') or 'N/A'} at {e.get('employer') or 'N/A'}"
                            f" . {e.get('start_date') or 'N/A'} to {e.get('end_date') or 'N/A'}"
                            f" . Tags: {', '.join(e.get('industry_tags', [])) or 'N/A'}"
                        )

                    st.write("**Summary**")
                    st.info(r.get("summary", ""))

                    st.markdown("**JSON**")
                    st.code(json.dumps(r, indent=2), language="json")

        st.markdown("---")
        st.subheader("Export batch JSON")
        batch = {
            "total_cvs": len(results),
            "processed_at": datetime.now().isoformat(),
            "cvs": results
        }
        data = json.dumps(batch, indent=2)
        st.download_button(
            label="📥 Download JSON",
            data=data,
            file_name=f"cvs_parsed_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json"
        )
    else:
        st.info("Process matched files to generate structured JSON outputs.")


if __name__ == "__main__":
    main()
