"""
Production-Ready CV Parser - Complete Application
Single file containing all functionality for CV/Resume parsing

Installation:
    pip install streamlit spacy pdfplumber python-docx pandas phonenumbers email-validator
    python -m spacy download en_core_web_lg

Usage:
    streamlit run cv_parser_complete.py
"""

import streamlit as st
import re
import json
import pandas as pd
from datetime import datetime
from typing import List, Dict, Optional
import io
import logging
from collections import Counter

# Try importing optional dependencies
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    
try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False
    
try:
    import phonenumbers
    PHONE_AVAILABLE = True
except ImportError:
    PHONE_AVAILABLE = False
    
try:
    from email_validator import validate_email, EmailNotValidError
    EMAIL_AVAILABLE = True
except ImportError:
    EMAIL_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ============================================================================
# CONFIGURATION
# ============================================================================

class Config:
    """Configuration for CV Parser"""
    
    # Supported formats
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc', '.txt']
    MAX_FILE_SIZE_MB = 10
    
    # Regex patterns
    PATTERNS = {
        'email': re.compile(
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            re.IGNORECASE
        ),
        'phone_uk': re.compile(
            r'(?:\+44\s?|0)(?:\d\s?){9,10}',
            re.IGNORECASE
        ),
        'phone_general': re.compile(
            r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            re.IGNORECASE
        ),
        'uk_postcode': re.compile(
            r'\b[A-Z]{1,2}\d{1,2}[A-Z]?\s?\d[A-Z]{2}\b',
            re.IGNORECASE
        ),
        'linkedin': re.compile(
            r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+/?',
            re.IGNORECASE
        ),
        'github': re.compile(
            r'(?:https?://)?(?:www\.)?github\.com/[\w-]+/?',
            re.IGNORECASE
        ),
        'date_range': re.compile(
            r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
            r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+'
            r'\d{4}\s*[-–—to]+\s*(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|'
            r'Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|'
            r'Dec(?:ember)?)\s+\d{4}|Present|Current)',
            re.IGNORECASE
        ),
        'year_range': re.compile(
            r'\b(19|20)\d{2}\s*[-–—to]+\s*(?:(19|20)\d{2}|Present|Current)\b',
            re.IGNORECASE
        ),
    }
    
    # Section headers
    SECTION_HEADERS = {
        'experience': ['experience', 'work history', 'employment', 'professional experience',
                      'work experience', 'career history'],
        'education': ['education', 'academic', 'qualifications'],
        'skills': ['skills', 'technical skills', 'competencies', 'expertise'],
        'certifications': ['certifications', 'certificates', 'licenses', 'accreditations'],
        'summary': ['summary', 'profile', 'objective', 'professional summary'],
    }
    
    # Keywords
    JOB_TITLES = ['manager', 'director', 'engineer', 'developer', 'analyst', 'consultant',
                  'specialist', 'supervisor', 'door supervisor', 'security officer', 'sia']
    
    CERTIFICATIONS_KEYWORDS = ['sia', 'door supervisor', 'security', 'first aid', 'cpr',
                               'cscs', 'nvq', 'certified', 'license', 'qualification']
    
    SKILLS_KEYWORDS = [
        # Technical
        'python', 'java', 'javascript', 'sql', 'html', 'css', 'react', 'aws', 'azure',
        # Security
        'sia', 'door supervisor', 'security guard', 'cctv', 'access control', 'patrol',
        # Soft skills
        'leadership', 'communication', 'teamwork', 'problem solving', 'customer service',
    ]

# ============================================================================
# DOCUMENT EXTRACTION
# ============================================================================

class DocumentExtractor:
    """Extract text from various document formats"""
    
    @staticmethod
    def extract_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF"""
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber not installed. Install with: pip install pdfplumber")
        
        text_content = []
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        return '\n\n'.join(text_content)
    
    @staticmethod
    def extract_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        doc = Document(io.BytesIO(file_content))
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_content.append(row_text)
        
        return '\n'.join(text_content)
    
    @staticmethod
    def extract_from_txt(file_content: bytes) -> str:
        """Extract text from TXT"""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return file_content.decode('utf-8', errors='ignore')
    
    @classmethod
    def extract_text(cls, file_content: bytes, filename: str) -> str:
        """Extract text based on file extension"""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return cls.extract_from_pdf(file_content)
        elif filename_lower.endswith(('.docx', '.doc')):
            return cls.extract_from_docx(file_content)
        elif filename_lower.endswith('.txt'):
            return cls.extract_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file format: {filename}")

# ============================================================================
# ENTITY EXTRACTION
# ============================================================================

class EntityExtractor:
    """Extract entities using regex patterns"""
    
    def __init__(self):
        self.patterns = Config.PATTERNS
    
    def extract_emails(self, text: str) -> List[str]:
        """Extract and validate emails"""
        emails = []
        potential_emails = self.patterns['email'].findall(text)
        
        for email in potential_emails:
            if EMAIL_AVAILABLE:
                try:
                    valid = validate_email(email, check_deliverability=False)
                    emails.append(valid.email)
                except EmailNotValidError:
                    continue
            else:
                emails.append(email)
        
        return list(dict.fromkeys(emails))
    
    def extract_phone_numbers(self, text: str) -> List[str]:
        """Extract phone numbers"""
        phone_numbers = []
        
        uk_phones = self.patterns['phone_uk'].findall(text)
        for phone in uk_phones:
            cleaned = re.sub(r'\s+', ' ', phone.strip())
            if cleaned not in phone_numbers:
                phone_numbers.append(cleaned)
        
        if not phone_numbers:
            general_phones = self.patterns['phone_general'].findall(text)
            for phone in general_phones:
                cleaned = re.sub(r'\s+', ' ', phone.strip())
                if cleaned not in phone_numbers and len(cleaned.replace(' ', '')) >= 10:
                    phone_numbers.append(cleaned)
        
        return phone_numbers[:3]
    
    def extract_postcodes(self, text: str) -> List[str]:
        """Extract UK postcodes"""
        postcodes = self.patterns['uk_postcode'].findall(text)
        valid_postcodes = []
        
        for postcode in postcodes:
            postcode = postcode.upper().strip().replace(' ', '')
            if len(postcode) >= 5:
                postcode = postcode[:-3] + ' ' + postcode[-3:]
                if postcode not in valid_postcodes:
                    valid_postcodes.append(postcode)
        
        return valid_postcodes
    
    def extract_urls(self, text: str) -> Dict[str, List[str]]:
        """Extract URLs"""
        return {
            'linkedin': list(dict.fromkeys(self.patterns['linkedin'].findall(text))),
            'github': list(dict.fromkeys(self.patterns['github'].findall(text)))
        }
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Divide CV into sections"""
        sections = {}
        lines = text.split('\n')
        current_section = 'header'
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            section_found = None
            
            for section_name, keywords in Config.SECTION_HEADERS.items():
                for keyword in keywords:
                    if keyword in line_lower and len(line_lower.split()) <= 4:
                        section_found = section_name
                        break
                if section_found:
                    break
            
            if section_found:
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = section_found
                current_content = []
            else:
                current_content.append(line)
        
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def extract_addresses(self, text: str, postcodes: List[str]) -> List[str]:
        """Extract addresses using postcodes"""
        addresses = []
        
        for postcode in postcodes:
            pattern = re.compile(
                r'([^\n]{0,150}' + re.escape(postcode) + r'[^\n]{0,20})',
                re.IGNORECASE
            )
            matches = pattern.findall(text)
            for match in matches:
                address = re.sub(r'\s+', ' ', match.strip())
                if len(address) > 10 and address not in addresses:
                    addresses.append(address)
        
        return addresses
    
    def extract_certifications(self, text: str) -> List[Dict[str, str]]:
        """Extract certifications"""
        certifications = []
        sections = self.extract_sections(text)
        cert_section = sections.get('certifications', text)
        
        lines = cert_section.split('\n')
        for line in lines:
            line_clean = line.strip()
            line_lower = line_clean.lower()
            
            if len(line_clean) < 3:
                continue
            
            for keyword in Config.CERTIFICATIONS_KEYWORDS:
                if keyword in line_lower:
                    date_match = re.search(r'\b(19|20)\d{2}\b', line_clean)
                    cert_dict = {
                        'name': line_clean,
                        'year': date_match.group(0) if date_match else None
                    }
                    if cert_dict not in certifications:
                        certifications.append(cert_dict)
                    break
        
        return certifications

# ============================================================================
# NLP EXTRACTION
# ============================================================================

class NLPExtractor:
    """Extract entities using NLP"""

    def __init__(self):
        if SPACY_AVAILABLE:
            try:
                self.nlp = spacy.load("en_core_web_lg")
            except OSError:
                try:
                    self.nlp = spacy.load("en_core_web_sm")
                except OSError:
                    self.nlp = None
                    logger.warning("No spaCy model found. Name extraction will be limited.")
        else:
            self.nlp = None

    def extract_name(self, text: str) -> Optional[str]:
        """Extract person's name (robust for CV headers)."""
        if not text:
            return None

        lines_raw = [ln.strip() for ln in text.splitlines() if ln.strip()]
        if not lines_raw:
            return None

        header_lines = lines_raw[:40]
        header_block = "\n".join(header_lines)

        bad_line_keywords = {
            "curriculum vitae", "cv", "resume", "résumé", "profile", "summary",
            "contact", "details", "personal details", "personal information",
            "email", "e-mail", "phone", "mobile", "tel", "telephone",
            "address", "linkedin", "github", "nationality", "dob", "date of birth",
            "postcode", "post code",
        }

        job_words = {
            "manager", "director", "engineer", "developer", "analyst", "consultant",
            "specialist", "supervisor", "officer", "coordinator", "assistant",
            "administrator", "executive", "lead", "technician",
        }

        def normalise_spaces(s: str) -> str:
            return re.sub(r"\s+", " ", s).strip()

        def strip_contact_noise(s: str) -> str:
            s = re.sub(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", " ", s)
            s = re.sub(r"(?:\+44\s?|0)(?:\d\s?){9,10}", " ", s, flags=re.I)
            s = re.sub(
                r"(?:https?://)?(?:www\.)?(linkedin\.com/in/\S+|github\.com/\S+)",
                " ",
                s,
                flags=re.I,
            )
            return normalise_spaces(s)

        def looks_like_bad_line(line: str) -> bool:
            l = line.lower()
            if any(k in l for k in bad_line_keywords):
                return True
            if len(line) > 80:
                return True
            digit_ratio = sum(ch.isdigit() for ch in line) / max(len(line), 1)
            if digit_ratio > 0.10:
                return True
            return False

        def title_case_if_all_caps(name: str) -> str:
            if name and name == name.upper():
                parts = name.split()
                return " ".join(p.capitalize() for p in parts)
            return name

        def is_plausible_name(name: str) -> bool:
            if not name:
                return False
            name = normalise_spaces(name)
            if len(name) < 3:
                return False

            if re.search(r"[^A-Za-z\s\-\.'’]", name):
                return False

            words = name.split()
            if len(words) < 2 or len(words) > 5:
                return False

            lw = [w.lower().strip(".") for w in words]
            if sum(w in job_words for w in lw) >= 1 and len(words) <= 3:
                return False

            joined = " ".join(lw)
            if any(k in joined for k in bad_line_keywords):
                return False

            return True

        def score_name(name: str, source: str, line_index: int) -> float:
            name = normalise_spaces(name)
            words = name.split()
            score = 0.0

            if source == "spacy":
                score += 3.0
            elif source == "regex":
                score += 2.0
            else:
                score += 1.0

            score += max(0.0, 4.0 - (line_index * 0.15))

            if len(words) == 2:
                score += 2.0
            elif len(words) == 3:
                score += 1.5
            elif len(words) == 4:
                score += 0.5
            else:
                score -= 0.5

            initials = sum(1 for w in words if re.fullmatch(r"[A-Za-z]\.?", w))
            if initials >= 2:
                score -= 1.0

            if name == name.upper():
                score -= 0.25

            for w in words:
                ww = w.strip(".")
                if len(ww) == 1:
                    continue
                if len(ww) < 2:
                    score -= 0.5

            return score

        candidates: List[Dict[str, object]] = []

        name_regex = re.compile(
            r"^(?P<name>"
            r"(?:[A-Za-z][A-Za-z'’\-]+|[A-Za-z]\.)"
            r"(?:\s+(?:[A-Za-z][A-Za-z'’\-]+|[A-Za-z]\.)){1,4}"
            r")$"
        )

        for i, line in enumerate(header_lines):
            if looks_like_bad_line(line):
                continue

            cleaned_line = strip_contact_noise(line)
            if not cleaned_line or looks_like_bad_line(cleaned_line):
                continue

            primary = re.split(r"[|,•·]", cleaned_line)[0].strip()
            primary = normalise_spaces(primary)

            m = name_regex.match(primary)
            if m:
                nm = title_case_if_all_caps(m.group("name"))
                nm = self._clean_name(nm)
                if is_plausible_name(nm):
                    candidates.append({"name": nm, "score": score_name(nm, "regex", i)})

            m2 = re.search(
                r"\bname\s*[:\-]\s*([A-Za-z][A-Za-z'’\-]+(?:\s+[A-Za-z][A-Za-z'’\-]+){1,4})\b",
                cleaned_line,
                re.I,
            )
            if m2:
                nm = title_case_if_all_caps(m2.group(1))
                nm = self._clean_name(nm)
                if is_plausible_name(nm):
                    candidates.append({"name": nm, "score": score_name(nm, "regex", i) + 1.0})

        if self.nlp:
            try:
                doc = self.nlp(header_block)
                for ent in doc.ents:
                    if ent.label_ != "PERSON":
                        continue

                    nm = strip_contact_noise(ent.text)
                    nm = title_case_if_all_caps(nm)
                    nm = self._clean_name(nm)

                    if is_plausible_name(nm):
                        line_index = 0
                        for idx, ln in enumerate(header_lines):
                            if ent.text in ln:
                                line_index = idx
                                break
                        candidates.append({"name": nm, "score": score_name(nm, "spacy", line_index)})
            except Exception:
                pass

        if not candidates:
            for i, line in enumerate(header_lines[:10]):
                if looks_like_bad_line(line):
                    continue
                line2 = strip_contact_noise(line)
                line2 = normalise_spaces(line2)
                m = re.match(
                    r"^([A-Za-z][A-Za-z'’\-]+(?:\s+[A-Za-z][A-Za-z'’\-]+){1,3})$",
                    line2,
                )
                if m:
                    nm = title_case_if_all_caps(m.group(1))
                    nm = self._clean_name(nm)
                    if is_plausible_name(nm):
                        return nm
            return None

        best_by_name: Dict[str, float] = {}
        for c in candidates:
            nm = str(c["name"])
            sc = float(c["score"])
            if nm not in best_by_name or sc > best_by_name[nm]:
                best_by_name[nm] = sc

        best_name = max(best_by_name.items(), key=lambda x: x[1])[0]
        return best_name

    def _clean_name(self, name: str) -> str:
        """Clean extracted name"""
        titles = {"mr", "mrs", "miss", "ms", "dr", "prof"}
        words = name.split()
        cleaned_words = [w for w in words if w.lower().strip(".") not in titles]
        return " ".join(cleaned_words).strip()

    def _is_valid_name(self, name: str) -> bool:
        """Validate name"""
        if not name or len(name) < 3:
            return False
        words = name.split()
        if len(words) < 2 or len(words) > 5:
            return False
        if any(char.isdigit() for char in name):
            return False
        return True

    def extract_experience(self, text: str) -> List[Dict]:
        """Extract professional experience"""
        experiences: List[Dict] = []
        sections = EntityExtractor().extract_sections(text)
        exp_text = sections.get("experience", text)

        date_pattern = re.compile(
            r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\b(?:19|20)\d{2}\b",
            re.IGNORECASE,
        )

        entries: List[str] = []
        current_entry: List[str] = []

        for line in exp_text.split("\n"):
            if date_pattern.search(line) and current_entry:
                entries.append("\n".join(current_entry))
                current_entry = [line]
            else:
                current_entry.append(line)

        if current_entry:
            entries.append("\n".join(current_entry))

        for entry in entries:
            if len(entry) < 20:
                continue

            lines = entry.split("\n")
            result = {
                "job_title": lines[0].strip() if lines else None,
                "company": None,
                "dates": None,
                "responsibilities": [],
            }

            date_match = Config.PATTERNS["date_range"].search(entry)
            if date_match:
                result["dates"] = date_match.group(0)

            for line in lines[1:]:
                line = re.sub(r"^[•\-\*\+○●]\s*", "", line.strip())
                if line and len(line) > 10:
                    result["responsibilities"].append(line)

            if result["job_title"]:
                experiences.append(result)

        return experiences

# ============================================================================
# SKILLS EXTRACTION
# ============================================================================

class SkillsExtractor:
    """Extract skills and keywords"""
    
    def __init__(self):
        self.skill_keywords = set(word.lower() for word in Config.SKILLS_KEYWORDS)
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from text"""
        text_lower = text.lower()
        found_skills = []
        
        for skill in self.skill_keywords:
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.append(skill)
        
        return list(dict.fromkeys(found_skills))
    
    def extract_keywords(self, text: str, top_n: int = 30) -> List[str]:
        """Extract top keywords"""
        text_lower = text.lower()
        
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'been'
        }
        
        words = re.findall(r'\b[a-z]{3,}\b', text_lower)
        filtered_words = [w for w in words if w not in stop_words]
        word_freq = Counter(filtered_words)
        
        # Boost important keywords
        boosted = {}
        for word, count in word_freq.items():
            boost = 3 if word in self.skill_keywords else 1
            boosted[word] = count * boost
        
        sorted_keywords = sorted(boosted.items(), key=lambda x: x[1], reverse=True)
        return [kw for kw, _ in sorted_keywords[:top_n]]
    
    def generate_summary(self, text: str, keywords: List[str], max_length: int = 150) -> str:
        """Generate summary"""
        sentences = re.split(r'[.!?]+', text)
        scored_sentences = []
        
        for sentence in sentences:
            if len(sentence.split()) < 5:
                continue
            score = sum(1 for kw in keywords[:15] if kw in sentence.lower())
            if score > 0:
                scored_sentences.append((sentence.strip(), score))
        
        scored_sentences.sort(key=lambda x: x[1], reverse=True)
        
        summary_parts = []
        word_count = 0
        
        for sentence, _ in scored_sentences[:3]:
            sentence_words = len(sentence.split())
            if word_count + sentence_words <= max_length:
                summary_parts.append(sentence)
                word_count += sentence_words
        
        return ' '.join(summary_parts) if summary_parts else ' '.join(text.split()[:max_length])

# ============================================================================
# CV PARSER
# ============================================================================

class CVParser:
    """Main CV parser"""
    
    def __init__(self):
        self.doc_extractor = DocumentExtractor()
        self.entity_extractor = EntityExtractor()
        self.nlp_extractor = NLPExtractor()
        self.skills_extractor = SkillsExtractor()
    
    def parse_cv(self, file_content: bytes, filename: str) -> Dict:
        """Parse CV and extract information"""
        try:
            # Extract text
            raw_text = self.doc_extractor.extract_text(file_content, filename)
            cleaned_text = ' '.join(raw_text.split())
            
            if not cleaned_text:
                raise ValueError("No text extracted")
            
            # Initialize result
            result = {
                'metadata': {
                    'filename': filename,
                    'processed_at': datetime.now().isoformat(),
                    'text_length': len(cleaned_text)
                },
                'personal_information': {},
                'contact_information': {},
                'professional_experience': [],
                'certifications_and_licenses': [],
                'skills': {},
                'keywords': [],
                'summary': ''
            }
            
            # Extract information
            name = self.nlp_extractor.extract_name(cleaned_text)
            result['personal_information']['full_name'] = name
            
            emails = self.entity_extractor.extract_emails(cleaned_text)
            phones = self.entity_extractor.extract_phone_numbers(cleaned_text)
            postcodes = self.entity_extractor.extract_postcodes(cleaned_text)
            addresses = self.entity_extractor.extract_addresses(cleaned_text, postcodes)
            urls = self.entity_extractor.extract_urls(cleaned_text)
            
            result['contact_information'] = {
                'email': emails[0] if emails else None,
                'all_emails': emails,
                'phone': phones[0] if phones else None,
                'all_phones': phones,
                'postcode': postcodes[0] if postcodes else None,
                'all_postcodes': postcodes,
                'address': addresses[0] if addresses else None,
                'linkedin': urls['linkedin'][0] if urls['linkedin'] else None,
                'github': urls['github'][0] if urls['github'] else None,
            }
            
            experiences = self.nlp_extractor.extract_experience(cleaned_text)
            result['professional_experience'] = experiences
            
            certifications = self.entity_extractor.extract_certifications(cleaned_text)
            result['certifications_and_licenses'] = certifications
            
            skills = self.skills_extractor.extract_skills(cleaned_text)
            result['skills'] = {'all_skills': skills}
            
            keywords = self.skills_extractor.extract_keywords(cleaned_text, top_n=30)
            result['keywords'] = keywords
            
            summary = self.skills_extractor.generate_summary(cleaned_text, keywords)
            result['summary'] = summary
            
            # Calculate insights
            result['insights'] = {
                'total_experience_entries': len(experiences),
                'total_certifications': len(certifications),
                'total_skills': len(skills),
                'has_email': bool(emails),
                'has_phone': bool(phones),
                'has_address': bool(addresses),
                'completeness_score': self._calculate_completeness(result)
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing CV: {str(e)}")
            raise
    
    def _calculate_completeness(self, data: Dict) -> int:
        """Calculate completeness score"""
        score = 0
        if data['personal_information'].get('full_name'):
            score += 20
        if data['contact_information'].get('email'):
            score += 15
        if data['contact_information'].get('phone'):
            score += 15
        if data['contact_information'].get('address'):
            score += 10
        if len(data['professional_experience']) > 0:
            score += 20
        if len(data['skills']['all_skills']) > 5:
            score += 10
        if len(data['certifications_and_licenses']) > 0:
            score += 10
        return min(score, 100)

# ============================================================================
# STREAMLIT APPLICATION
# ============================================================================

def main():
    """Main Streamlit application"""
    
    st.set_page_config(
        page_title="CV Parser Pro",
        page_icon="📄",
        layout="wide"
    )
    
    # Custom CSS
    st.markdown("""
        <style>
        .main-header {font-size: 3rem; font-weight: bold; color: #1f77b4; text-align: center;}
        .metric-card {background-color: #ffffff; padding: 1.5rem; border-radius: 0.5rem; 
                      box-shadow: 0 2px 4px rgba(0,0,0,0.1); text-align: center;}
        .metric-value {font-size: 2.5rem; font-weight: bold; color: #1f77b4;}
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="main-header">📄 CV Parser Pro</div>', unsafe_allow_html=True)
    st.markdown("### Production-Ready Resume Information Extraction System")
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Configuration")
        st.info("Supported formats: PDF, DOCX, TXT")
        
        st.markdown("### Features")
        features = [
            "✅ Name & contact extraction",
            "✅ Experience parsing",
            "✅ Skills identification",
            "✅ Certifications detection",
            "✅ Keyword extraction",
            "✅ JSON export"
        ]
        for feature in features:
            st.markdown(feature)
    
    # Main tabs
    tab1, tab2 = st.tabs(["📤 Upload & Process", "📊 Results"])
    
    with tab1:
        st.markdown("## Upload CV Files")
        
        uploaded_files = st.file_uploader(
            "Choose CV files",
            type=['pdf', 'docx', 'doc', 'txt'],
            accept_multiple_files=True
        )
        
        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} file(s) uploaded")
            
            if st.button("🚀 Process CVs", type="primary"):
                with st.spinner("Processing CVs..."):
                    try:
                        parser = CVParser()
                        results = []
                        
                        for file in uploaded_files:
                            result = parser.parse_cv(file.getvalue(), file.name)
                            results.append(result)
                        
                        st.session_state['parsed_cvs'] = results
                        st.success("✅ Processing complete!")
                        st.balloons()
                        
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")
    
    with tab2:
        st.markdown("## Extraction Results")
        
        if 'parsed_cvs' in st.session_state:
            parsed_cvs = st.session_state['parsed_cvs']
            
            # Metrics
            col1, col2, col3, col4 = st.columns(4)
            
            total = len(parsed_cvs)
            avg_completeness = sum(cv['insights']['completeness_score'] for cv in parsed_cvs) / total
            total_skills = sum(cv['insights']['total_skills'] for cv in parsed_cvs)
            
            with col1:
                st.metric("Total CVs", total)
            with col2:
                st.metric("Avg Completeness", f"{avg_completeness:.0f}%")
            with col3:
                st.metric("Total Skills", total_skills)
            with col4:
                st.metric("Success Rate", "100%")
            
            st.markdown("---")
            
            # Detailed results
            for idx, cv in enumerate(parsed_cvs, 1):
                with st.expander(f"CV {idx}: {cv['metadata']['filename']}", expanded=(idx == 1)):
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.write("**Name:**", cv['personal_information'].get('full_name', 'Not found'))
                        st.write("**Email:**", cv['contact_information'].get('email', 'Not found'))
                        st.write("**Phone:**", cv['contact_information'].get('phone', 'Not found'))
                    with col2:
                        st.write("**Postcode:**", cv['contact_information'].get('postcode', 'Not found'))
                        st.write("**LinkedIn:**", cv['contact_information'].get('linkedin', 'Not found'))
                        completeness = cv['insights']['completeness_score']
                        st.write(f"**Completeness:** {completeness}%")
                        st.progress(completeness / 100)
                    
                    if cv['professional_experience']:
                        st.markdown("**Experience:**")
                        for exp in cv['professional_experience'][:3]:
                            st.write(f"• {exp.get('job_title', 'N/A')} - {exp.get('dates', 'N/A')}")
                    
                    if cv['certifications_and_licenses']:
                        st.markdown("**Certifications:**")
                        for cert in cv['certifications_and_licenses'][:5]:
                            st.write(f"• {cert['name']}")
                    
                    if cv['skills']['all_skills']:
                        st.markdown("**Skills:**")
                        st.write(", ".join(cv['skills']['all_skills'][:15]))
                    
                    if cv['keywords']:
                        st.markdown("**Keywords:**")
                        st.write(", ".join(cv['keywords'][:15]))
                    
                    st.markdown("**Summary:**")
                    st.info(cv['summary'])
            
            # Export
            st.markdown("---")
            st.markdown("### 💾 Export")
            
            json_output = json.dumps({
                'total_cvs': len(parsed_cvs),
                'processed_at': datetime.now().isoformat(),
                'cvs': parsed_cvs
            }, indent=2)
            
            st.download_button(
                label="📥 Download JSON",
                data=json_output,
                file_name=f"cvs_parsed_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )
            st.code(json_output, language="json")
        else:
            st.info("👆 Upload and process CVs to see results")

if __name__ == "__main__":
    main()