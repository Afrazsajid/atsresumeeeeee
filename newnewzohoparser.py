# app.py
# Streamlit CV Sync + Hybrid Search (Security, Warehouse, Hospitality)
# Stores parsed candidates into sector JSON files and supports advanced query matching.

import os
import io
import re
import json
import math
import time
import logging
import traceback
from dataclasses import dataclass, field
from pathlib import Path
from datetime import datetime, date
from typing import Dict, List, Optional, Tuple, Any
from urllib.parse import urlparse, parse_qs

import streamlit as st
import pandas as pd
import numpy as np
import requests

# Optional dependencies
try:
    import pdfplumber
    PDF_AVAILABLE = True
except Exception:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

import spacy
SPACY_AVAILABLE = True


try:
    import phonenumbers
    PHONE_AVAILABLE = True
except Exception:
    PHONE_AVAILABLE = False

try:
    from email_validator import validate_email, EmailNotValidError
    EMAIL_AVAILABLE = True
except Exception:
    EMAIL_AVAILABLE = False

try:
    from rapidfuzz import fuzz
    RAPIDFUZZ_AVAILABLE = True
except Exception:
    RAPIDFUZZ_AVAILABLE = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    SKLEARN_AVAILABLE = True
except Exception:
    SKLEARN_AVAILABLE = False

# OCR optional
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False


# ---------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------
logger = logging.getLogger("cv_parser")
logger.setLevel(logging.INFO)
_handler = logging.StreamHandler()
_handler.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
if not logger.handlers:
    logger.addHandler(_handler)


# ---------------------------------------------------------------------
# Resume filename pattern
# ---------------------------------------------------------------------
RESUME_STEM_RE = re.compile(r"^(resume|cv)(?:[-_ ]|$)", re.IGNORECASE)


def is_resume_filename(name: str, supported_exts: Tuple[str, ...], exact_filename: str = "") -> bool:
    if not name:
        return False
    if exact_filename and name.strip().lower() != exact_filename.strip().lower():
        return False
    p = Path(name)
    if p.suffix.lower() not in supported_exts:
        return False
    return bool(RESUME_STEM_RE.match(p.stem))


# ---------------------------------------------------------------------
# Zoho WorkDrive Client
# ---------------------------------------------------------------------
class ZohoWorkDriveClient:
    """
    Required secrets in .streamlit/secrets.toml:
      ZOHO_CLIENT_ID
      ZOHO_CLIENT_SECRET
      ZOHO_REFRESH_TOKEN

    Optional:
      ZOHO_ACCOUNTS_DOMAIN
      ZOHO_API_BASE
      ZOHO_WORKDRIVE_WEB_BASE
    """

    def __init__(self):
        self.client_id = st.secrets.get("ZOHO_CLIENT_ID", "")
        self.client_secret = st.secrets.get("ZOHO_CLIENT_SECRET", "")
        self.refresh_token = st.secrets.get("ZOHO_REFRESH_TOKEN", "")

        self.accounts_domain = st.secrets.get("ZOHO_ACCOUNTS_DOMAIN", "https://accounts.zoho.com").rstrip("/")
        self.api_base = st.secrets.get("ZOHO_API_BASE", "https://www.zohoapis.com").rstrip("/")
        self.workdrive_web_base = st.secrets.get("ZOHO_WORKDRIVE_WEB_BASE", "https://workdrive.zoho.com").rstrip("/")

        if not (self.client_id and self.client_secret and self.refresh_token):
            raise ValueError("Missing Zoho secrets. Add ZOHO_CLIENT_ID, ZOHO_CLIENT_SECRET, ZOHO_REFRESH_TOKEN to secrets.toml")

        self.workdrive_api_base = f"{self.api_base}/workdrive/api/v1"

    def _get_cached_token(self) -> Optional[str]:
        tok = st.session_state.get("zoho_access_token")
        exp = st.session_state.get("zoho_access_token_expires_at", 0)
        if tok and time.time() < exp - 30:
            return tok
        return None

    def _cache_token(self, token: str, expires_in_sec: int):
        st.session_state["zoho_access_token"] = token
        st.session_state["zoho_access_token_expires_at"] = time.time() + int(expires_in_sec)

    def refresh_access_token(self) -> str:
        cached = self._get_cached_token()
        if cached:
            return cached

        url = f"{self.accounts_domain}/oauth/v2/token"
        data = {
            "grant_type": "refresh_token",
            "client_id": self.client_id,
            "client_secret": self.client_secret,
            "refresh_token": self.refresh_token,
        }
        r = requests.post(url, data=data, timeout=30)
        if r.status_code != 200:
            raise RuntimeError(f"Token refresh failed ({r.status_code}). {r.text}")

        j = r.json()
        token = j.get("access_token")
        expires = int(j.get("expires_in_sec") or j.get("expires_in") or 3600)
        if not token:
            raise RuntimeError(f"Token refresh response missing access_token. {j}")

        self._cache_token(token, expires)
        return token

    def _headers(self) -> Dict[str, str]:
        return {
            "Authorization": f"Zoho-oauthtoken {self.refresh_access_token()}",
            "Accept": "application/vnd.api+json",
        }

    def _request(self, method: str, url: str, **kwargs) -> requests.Response:
        r = requests.request(method, url, headers=self._headers(), timeout=kwargs.pop("timeout", 60), **kwargs)
        if r.status_code == 401:
            st.session_state.pop("zoho_access_token", None)
            st.session_state.pop("zoho_access_token_expires_at", None)
            r = requests.request(method, url, headers=self._headers(), timeout=kwargs.pop("timeout", 60), **kwargs)
        return r

    @staticmethod
    def extract_resource_id(folder_link_or_id: str) -> str:
        s = (folder_link_or_id or "").strip()
        if not s:
            raise ValueError("Folder link or id is empty.")

        if re.fullmatch(r"[A-Za-z0-9]{8,}", s):
            return s

        u = urlparse(s)
        parts = [p for p in u.path.split("/") if p]
        for p in reversed(parts):
            if re.fullmatch(r"[A-Za-z0-9]{8,}", p):
                return p

        qs = parse_qs(u.query or "")
        for key in ["id", "res_id", "folder_id", "resource_id"]:
            if key in qs and qs[key]:
                val = qs[key][0]
                if re.fullmatch(r"[A-Za-z0-9]{8,}", val):
                    return val

        m = re.search(r"(?:folder|folders|file|files|id|res_id)=([A-Za-z0-9]{8,})", s, flags=re.I)
        if m:
            return m.group(1)

        raise ValueError("Could not detect folder id from link. Paste the folder id directly.")

    @staticmethod
    def _item_name(it: Dict[str, Any]) -> str:
        attrs = it.get("attributes") or {}
        return (attrs.get("name") or it.get("name") or "").strip()

    @staticmethod
    def _item_id(it: Dict[str, Any]) -> str:
        return (it.get("id") or "").strip()

    @staticmethod
    def _is_folder_item(it: Dict[str, Any]) -> bool:
        t = (it.get("type") or "").lower()
        attrs = it.get("attributes") or {}
        if t == "folders":
            return True
        if attrs.get("is_folder") is True:
            return True
        if (attrs.get("type") or "").lower() == "folder":
            return True
        if (attrs.get("mime_type") or "").lower() in {"folder", "application/vnd.zoho.folder"}:
            return True
        if (attrs.get("file_type") or "").lower() == "folder":
            return True
        return False

    @staticmethod
    def _is_file_item(it: Dict[str, Any]) -> bool:
        if ZohoWorkDriveClient._is_folder_item(it):
            return False
        t = (it.get("type") or "").lower()
        if t == "files":
            return True
        attrs = it.get("attributes") or {}
        if attrs.get("is_folder") is False:
            return True
        return bool(ZohoWorkDriveClient._item_name(it))

    def _list_children_one_page(self, folder_id: str, offset: int, limit: int) -> Tuple[List[Dict[str, Any]], bool]:
        endpoints = [
            f"{self.workdrive_api_base}/files/{folder_id}/files",
            f"{self.workdrive_api_base}/folders/{folder_id}/files",
            f"{self.workdrive_api_base}/folders/{folder_id}/children",
            f"{self.workdrive_api_base}/files/{folder_id}/children",
        ]

        last_err = None
        for url in endpoints:
            params = {"filter[type]": "all", "page[limit]": limit, "page[offset]": offset}
            try:
                r = self._request("GET", url, params=params, timeout=60)
                if r.status_code != 200:
                    last_err = f"{url} -> {r.status_code} {r.text[:400]}"
                    continue

                payload = r.json() if r.text else {}
                items = payload.get("data", []) or []
                included = payload.get("included", []) or []
                if included:
                    items = list(items) + list(included)
                items = [x for x in items if isinstance(x, dict)]

                has_more = False
                links = payload.get("links") or {}
                if isinstance(links, dict) and links.get("next"):
                    has_more = True
                else:
                    has_more = len(payload.get("data", []) or []) >= limit

                return items, has_more
            except Exception as e:
                last_err = f"{url} -> {e}"
                continue

        raise RuntimeError(f"Could not list folder children. Last error: {last_err}")

    def list_resume_files_recursive(
        self,
        root_folder_id: str,
        include_subfolders: bool,
        supported_exts: Tuple[str, ...],
        exact_filename: str = "",
        max_items_safety: int = 20000
    ) -> List[Dict[str, Any]]:
        results: List[Dict[str, Any]] = []
        visited = set()

        queue: List[Tuple[str, str]] = [(root_folder_id, "")]
        scanned = 0

        while queue:
            folder_id, prefix = queue.pop(0)
            if folder_id in visited:
                continue
            visited.add(folder_id)

            offset = 0
            limit = 50
            loops = 0

            while True:
                loops += 1
                if loops > 5000:
                    raise RuntimeError("Safety stop. Too many pagination loops while listing WorkDrive folder contents.")

                items, has_more = self._list_children_one_page(folder_id, offset=offset, limit=limit)
                scanned += len(items)
                if scanned > max_items_safety:
                    raise RuntimeError("Safety stop. Too many items scanned. Narrow folder or increase max_items_safety.")

                for it in items:
                    name = self._item_name(it)
                    it_id = self._item_id(it)
                    if not it_id or not name:
                        continue

                    if self._is_folder_item(it):
                        if include_subfolders:
                            new_prefix = f"{prefix}{name}/"
                            queue.append((it_id, new_prefix))
                        continue

                    if self._is_file_item(it):
                        if is_resume_filename(name, supported_exts=supported_exts, exact_filename=exact_filename):
                            results.append({
                                "id": it_id,
                                "name": name,
                                "virtual_path": f"{prefix}{name}".strip(),
                            })

                if not has_more:
                    break
                offset += limit

        results.sort(key=lambda x: (x.get("virtual_path") or x.get("name") or "").lower())
        return results

    def download_file_bytes(self, file_id: str) -> bytes:
        url1 = f"{self.workdrive_web_base}/api/v1/download/{file_id}"
        r = self._request("GET", url1, timeout=120, allow_redirects=True)
        if r.status_code == 200 and r.content:
            return r.content

        url2 = "https://download.zoho.com/v1/workdrive/download/" + file_id
        r2 = self._request("GET", url2, timeout=120, allow_redirects=True)
        if r2.status_code == 200 and r2.content:
            return r2.content

        url3 = f"{self.workdrive_api_base}/download/{file_id}"
        r3 = self._request("GET", url3, timeout=120, allow_redirects=True)
        if r3.status_code == 200 and r3.content:
            return r3.content

        raise RuntimeError(
            f"Download failed for file_id={file_id}. "
            f"status1={r.status_code}, status2={r2.status_code}, status3={r3.status_code}"
        )

    def file_web_url(self, file_id: str) -> str:
        # This usually works for most tenants. If your tenant uses a different pattern, change here.
        return f"{self.workdrive_web_base}/file/{file_id}"


# ---------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------
MONTHS = {
    "jan": 1, "january": 1,
    "feb": 2, "february": 2,
    "mar": 3, "march": 3,
    "apr": 4, "april": 4,
    "may": 5,
    "jun": 6, "june": 6,
    "jul": 7, "july": 7,
    "aug": 8, "august": 8,
    "sep": 9, "sept": 9, "september": 9,
    "oct": 10, "october": 10,
    "nov": 11, "november": 11,
    "dec": 12, "december": 12,
}

@dataclass
class AppConfig:
    supported_exts: Tuple[str, ...] = (".pdf", ".docx")
    max_file_mb: int = 25

    spacy_model_preference: Tuple[str, ...] = ("en_core_web_lg", "en_core_web_md", "en_core_web_sm")

    min_extracted_alpha_chars_for_pdf: int = 400
    drop_repeated_header_footer_ratio: float = 0.60

    enable_ocr_fallback: bool = True
    ocr_lang: str = "eng"

    embedding_similarity_threshold: float = 0.70
    keep_unmatched_skill_phrases: bool = True

    max_workers: int = 4

    section_aliases: Dict[str, List[str]] = field(default_factory=lambda: {
        "summary": ["summary", "profile", "professional summary", "about", "objective"],
        "skills": ["skills", "technical skills", "core skills", "key skills", "competencies", "expertise"],
        "experience": ["experience", "work experience", "employment", "career history", "work history", "professional experience"],
        "education": ["education", "qualifications", "academic background"],
        "certifications": ["certifications", "certificates", "licences", "licenses", "training", "accreditations"],
        "contact": ["contact", "contact details", "personal details", "personal information"],
    })

    industry_keywords: Dict[str, List[str]] = field(default_factory=lambda: {
        "security": [
            "sia", "door supervisor", "security officer", "cctv", "steward", "event security", "patrol",
            "access control", "crowd", "incident", "search procedure", "id check", "safeguarding"
        ],
        "warehouse": [
            "picker", "packer", "picking", "packing", "rf scanner", "scan gun", "forklift", "reach truck",
            "loading", "unloading", "inventory", "stock", "dispatch", "goods in", "goods out"
        ],
        "hospitality": [
            "housekeeping", "room attendant", "waiter", "waitress", "bar", "bartender", "front of house",
            "concierge", "porter", "linen porter", "service standards"
        ]
    })

    email_re: re.Pattern = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", re.I)
    uk_postcode_re: re.Pattern = re.compile(r"\b([A-Z]{1,2}\d{1,2}[A-Z]?)\s*(\d[A-Z]{2})\b", re.I)

    date_range_re: re.Pattern = re.compile(
        r"(?P<start>(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
        r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t)?(?:ember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}|\b(19|20)\d{2}\b)"
        r"\s*(?:-+|–|to)\s*"
        r"(?P<end>(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
        r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t)?(?:ember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}|\b(19|20)\d{2}\b|present|current|now)"
        , re.I
    )

    bullet_re: re.Pattern = re.compile(r"^\s*([•\-\*\u2022\u25CF\u25CB\u25AA\u00B7]+|\d+\.)\s+")

CONFIG = AppConfig()


# ---------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------
def dedupe_preserve_order(items: List[str]) -> List[str]:
    seen = set()
    out = []
    for x in items:
        k = x.strip()
        if not k:
            continue
        if k.lower() in seen:
            continue
        seen.add(k.lower())
        out.append(k)
    return out


def normalise_whitespace_keep_lines(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = []
    for ln in text.split("\n"):
        ln2 = re.sub(r"[ \t]+", " ", ln).strip()
        lines.append(ln2)
    out = "\n".join(lines)
    out = re.sub(r"\n{3,}", "\n\n", out).strip()
    return out


def safe_clip01(x: float) -> float:
    return float(max(0.0, min(1.0, x)))


def cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
    if a is None or b is None:
        return 0.0
    na = np.linalg.norm(a)
    nb = np.linalg.norm(b)
    if na == 0.0 or nb == 0.0:
        return 0.0
    return float(np.dot(a, b) / (na * nb))


def looks_like_heading(line: str) -> bool:
    if not line:
        return False
    l = line.strip()
    if len(l) < 3:
        return False
    words = l.split()
    if len(words) > 6:
        return False
    if l.endswith(":"):
        return True
    if l.upper() == l and any(c.isalpha() for c in l) and len(l) <= 40:
        return True
    return False


def normalise_postcode(m1: str, m2: str) -> str:
    return f"{m1.upper()} {m2.upper()}"


def parse_month_year(s: str) -> Tuple[Optional[date], float]:
    s0 = s.strip().lower()
    if re.fullmatch(r"(19|20)\d{2}", s0):
        try:
            y = int(s0)
            return date(y, 1, 1), 0.70
        except Exception:
            return None, 0.0

    m = re.match(r"^(?P<mon>[a-z]{3,9})\s+(?P<yr>(19|20)\d{2})$", s0)
    if not m:
        return None, 0.0
    mon = m.group("mon")
    yr = int(m.group("yr"))
    mon_num = MONTHS.get(mon[:3], None) if mon not in MONTHS else MONTHS[mon]
    if not mon_num:
        mon_num = MONTHS.get(mon, None)
    if not mon_num:
        return None, 0.0
    return date(yr, mon_num, 1), 0.90


def date_to_yyyy_mm(d: date) -> str:
    return f"{d.year:04d}-{d.month:02d}"


def months_between(a: date, b: date) -> int:
    return (b.year - a.year) * 12 + (b.month - a.month)


# ---------------------------------------------------------------------
# spaCy loader
# ---------------------------------------------------------------------
@st.cache_resource(show_spinner=False)
def load_spacy_model():
    if not SPACY_AVAILABLE:
        return None
    for name in CONFIG.spacy_model_preference:
        try:
            nlp = spacy.load(name)
            return nlp
        except Exception:
            continue
    return None


# ---------------------------------------------------------------------
# Skill catalogue (kept compact)
# ---------------------------------------------------------------------
DEFAULT_SKILL_CATALOGUE: Dict[str, Any] = {
    "canonical": [
        "SIA",
        "SIA Door Supervisor",
        "SIA CCTV",
        "CCTV monitoring",
        "Access control",
        "Incident reporting",
        "Conflict management",
        "Crowd management",
        "ID verification",
        "Search procedures",
        "Emergency response procedures",
        "Patrolling",
        "Radio communication",
        "Safeguarding",
        "Customer service",
        "Health and safety compliance",
        "GDPR and data protection compliance",
        "Picking and packing",
        "Order picking",
        "RF scanning",
        "Stock control",
        "Inventory management",
        "Loading and unloading",
        "Forklift driving",
        "Reach truck",
        "Manual handling",
        "PPE compliance",
        "Housekeeping",
        "Room attendant",
        "Front of house",
        "Food hygiene",
        "Bar service",
        "Microsoft Excel",
        "Microsoft Word",
        "Team leadership",
        "Communication",
        "Problem solving",
        "Time management",
    ],
    "variants": {
        "SIA Door Supervisor": ["door supervisor", "sia ds", "ds licence", "ds license", "door sup", "sia door super visor"],
        "SIA CCTV": ["cctv licence", "cctv license", "sia cctv"],
        "CCTV monitoring": ["cctv", "monitoring cctv", "cctv operator"],
        "Picking and packing": ["picker", "packer", "picking", "packing", "pick and pack"],
        "RF scanning": ["rf scanner", "rf scan", "scan gun", "handheld scanner"],
        "Forklift driving": ["forklift", "forklift driver", "counterbalance"],
        "Reach truck": ["reach", "reach driver", "reach forklift"],
        "Housekeeping": ["house keeper", "housekeeper", "public area cleaner"],
    },
    "abbreviations": {
        "DS": ["SIA Door Supervisor"],
        "CCTV": ["CCTV monitoring", "SIA CCTV"],
        "GDPR": ["GDPR and data protection compliance"],
        "DBS": ["DBS check"],
        "HSE": ["Health and safety compliance"],
    }
}


class SkillCatalogue:
    def __init__(self, data: Dict[str, Any], nlp=None):
        self.data = data
        self.nlp = nlp
        self.canonical: List[str] = data.get("canonical", [])
        self.variants: Dict[str, List[str]] = data.get("variants", {})
        self.abbreviations: Dict[str, List[str]] = data.get("abbreviations", {})

        self.variant_to_canonical: Dict[str, str] = {}
        for canon, vars_ in self.variants.items():
            for v in vars_:
                self.variant_to_canonical[v.lower().strip()] = canon

        self._canon_vectors: Dict[str, np.ndarray] = {}
        if self.nlp and getattr(self.nlp.vocab, "vectors", None) and self.nlp.vocab.vectors.size > 0:
            for canon in self.canonical:
                self._canon_vectors[canon] = self.nlp(canon).vector

    @staticmethod
    def load_from_path(path: Optional[str], nlp=None) -> "SkillCatalogue":
        if path:
            p = Path(path)
            if p.exists() and p.is_file():
                with p.open("r", encoding="utf-8") as f:
                    data = json.load(f)
                return SkillCatalogue(data, nlp=nlp)
        return SkillCatalogue(DEFAULT_SKILL_CATALOGUE, nlp=nlp)

    def canonicalise(self, text: str, context_text: str = "") -> Tuple[str, float, str]:
        t = text.strip()
        if not t:
            return "", 0.0, "passthrough"

        low = t.lower()

        for canon in self.canonical:
            if low == canon.lower():
                return canon, 0.98, "exact"

        if low in self.variant_to_canonical:
            return self.variant_to_canonical[low], 0.95, "variant"

        up = re.sub(r"[^A-Za-z]", "", t).upper()
        if up in self.abbreviations:
            candidates = self.abbreviations[up]
            if len(candidates) == 1:
                return candidates[0], 0.88, "abbrev"
            return candidates[0], 0.70, "abbrev"

        if self.nlp and self._canon_vectors:
            v = self.nlp(t).vector
            best = ("", 0.0)
            for canon, cv in self._canon_vectors.items():
                sim = cosine_sim(v, cv)
                if sim > best[1]:
                    best = (canon, sim)
            if best[1] >= CONFIG.embedding_similarity_threshold:
                conf = safe_clip01((best[1] - CONFIG.embedding_similarity_threshold) / (1.0 - CONFIG.embedding_similarity_threshold))
                conf = 0.60 + 0.35 * conf
                return best[0], conf, "embed"

        return self._title_like(t), 0.55, "passthrough"

    @staticmethod
    def _title_like(s: str) -> str:
        s2 = re.sub(r"\s+", " ", s).strip()
        if not s2:
            return s2
        if s2.isupper():
            return " ".join(w.capitalize() for w in s2.split())
        return s2


# ---------------------------------------------------------------------
# Document extraction
# ---------------------------------------------------------------------
class DocumentExtractor:
    def extract_text(self, file_bytes: bytes, filename: str) -> Dict[str, str]:
        name = filename.lower()
        if name.endswith(".pdf"):
            raw = self._extract_pdf(file_bytes)
        elif name.endswith(".docx"):
            raw = self._extract_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {filename}")

        clean = normalise_whitespace_keep_lines(raw)
        return {"raw_text": raw, "clean_text": clean}

    def _extract_docx(self, file_bytes: bytes) -> str:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. pip install python-docx")

        doc = Document(io.BytesIO(file_bytes))
        parts = []

        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)

        for tbl in doc.tables:
            for row in tbl.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    parts.append(row_text)

        return "\n".join(parts).strip()

    def _extract_pdf(self, file_bytes: bytes) -> str:
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber is not installed. pip install pdfplumber")

        pages_text = []
        page_lines_list = []

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                txt = ""
                try:
                    txt = page.extract_text(layout=True, x_tolerance=2, y_tolerance=2) or ""
                except Exception:
                    txt = ""

                if not txt:
                    txt = self._reconstruct_from_words(page) or ""

                pages_text.append(txt)
                lines = [ln.strip() for ln in txt.split("\n") if ln.strip()]
                page_lines_list.append(lines)

        cleaned_pages = self._drop_repeated_header_footer(page_lines_list)
        joined = "\n\n".join("\n".join(lines) for lines in cleaned_pages).strip()

        alpha_chars = sum(ch.isalpha() for ch in joined)
        if CONFIG.enable_ocr_fallback and alpha_chars < CONFIG.min_extracted_alpha_chars_for_pdf:
            ocr_text = self._ocr_pdf(file_bytes)
            if ocr_text and sum(ch.isalpha() for ch in ocr_text) > alpha_chars:
                return ocr_text

        return joined

    @staticmethod
    def _reconstruct_from_words(page) -> str:
        try:
            words = page.extract_words(use_text_flow=True, keep_blank_chars=False)
            if not words:
                return ""
            words_sorted = sorted(words, key=lambda w: (round(w["top"], 1), w["x0"]))
            lines = []
            current_top = None
            current = []
            for w in words_sorted:
                top = round(w["top"], 1)
                if current_top is None:
                    current_top = top
                if abs(top - current_top) > 2.0:
                    lines.append(" ".join(current).strip())
                    current = [w["text"]]
                    current_top = top
                else:
                    current.append(w["text"])
            if current:
                lines.append(" ".join(current).strip())
            return "\n".join([ln for ln in lines if ln.strip()])
        except Exception:
            return ""

    @staticmethod
    def _drop_repeated_header_footer(page_lines_list: List[List[str]]) -> List[List[str]]:
        if not page_lines_list:
            return page_lines_list

        n_pages = len(page_lines_list)
        if n_pages < 2:
            return page_lines_list

        freq: Dict[str, int] = {}
        for lines in page_lines_list:
            unique = set(lines)
            for ln in unique:
                if len(ln) <= 80:
                    freq[ln] = freq.get(ln, 0) + 1

        threshold = max(2, int(math.ceil(n_pages * CONFIG.drop_repeated_header_footer_ratio)))
        to_remove = {ln for ln, c in freq.items() if c >= threshold}

        cleaned = []
        for lines in page_lines_list:
            out = []
            for ln in lines:
                if ln in to_remove:
                    continue
                out.append(ln)
            cleaned.append(out)
        return cleaned

    @staticmethod
    def _ocr_pdf(file_bytes: bytes) -> str:
        if not OCR_AVAILABLE:
            logger.info("OCR fallback requested but OCR dependencies are not installed.")
            return ""

        try:
            images = convert_from_bytes(file_bytes)
            texts = []
            for img in images:
                t = pytesseract.image_to_string(img, lang=CONFIG.ocr_lang)
                if t:
                    texts.append(t)
            return "\n\n".join(texts).strip()
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return ""


# ---------------------------------------------------------------------
# Section parsing
# ---------------------------------------------------------------------
class SectionParser:
    def __init__(self, config: AppConfig):
        self.config = config
        self.header_to_section: Dict[str, str] = {}
        for section, aliases in config.section_aliases.items():
            for a in aliases:
                self.header_to_section[a.lower()] = section

    def split_sections(self, text: str) -> Dict[str, List[str]]:
        lines = [ln.strip() for ln in text.split("\n")]
        sections: Dict[str, List[str]] = {}
        current = "header"
        buf: List[str] = []

        def flush():
            nonlocal buf, current
            if buf:
                sections.setdefault(current, [])
                sections[current].extend([x for x in buf if x.strip()])
            buf = []

        for ln in lines:
            lnl = ln.lower().strip(": ").strip()
            section_found = None

            if looks_like_heading(ln):
                lnl2 = lnl.replace("&", "and")
                for key, sec in self.header_to_section.items():
                    if lnl2 == key:
                        section_found = sec
                        break
                    if lnl2 in key or key in lnl2:
                        if len(lnl2.split()) <= 4:
                            section_found = sec
                            break

            if section_found:
                flush()
                current = section_found
            else:
                buf.append(ln)

        flush()
        return sections


# ---------------------------------------------------------------------
# Contact extraction
# ---------------------------------------------------------------------
class ContactExtractor:
    def __init__(self, config: AppConfig):
        self.config = config

    def extract_emails(self, text: str) -> List[str]:
        found = self.config.email_re.findall(text)
        out = []
        for e in found:
            if EMAIL_AVAILABLE:
                try:
                    v = validate_email(e, check_deliverability=False)
                    out.append(v.email)
                except EmailNotValidError:
                    continue
            else:
                out.append(e)
        return dedupe_preserve_order(out)

    def extract_uk_phones(self, text: str) -> List[str]:
        candidates = []
        raw = re.findall(r"(\+44\s?\d[\d\s]{8,13}|\b0\d[\d\s]{8,13})", text)
        raw = [re.sub(r"\s+", " ", r).strip() for r in raw]
        raw = dedupe_preserve_order(raw)

        if PHONE_AVAILABLE:
            for r in raw:
                try:
                    num = phonenumbers.parse(r, "GB")
                    if phonenumbers.is_possible_number(num) and phonenumbers.is_valid_number(num):
                        fmt = phonenumbers.format_number(num, phonenumbers.PhoneNumberFormat.E164)
                        candidates.append(fmt)
                except Exception:
                    continue
            candidates = dedupe_preserve_order(candidates)
            if candidates:
                return candidates[:3]

        cleaned = []
        for r in raw:
            digits = re.sub(r"\D", "", r)
            if len(digits) >= 10:
                cleaned.append(r)
        return dedupe_preserve_order(cleaned)[:3]

    def extract_postcodes(self, text: str) -> List[str]:
        matches = self.config.uk_postcode_re.findall(text)
        out = []
        for m1, m2 in matches:
            pc = normalise_postcode(m1, m2)
            if pc not in out:
                out.append(pc)
        return out

    def extract_address(self, text: str, postcodes: List[str]) -> Tuple[Optional[str], float, str]:
        if not postcodes:
            return None, 0.0, "none"

        lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
        best = ("", 0.0, "postcode_window")

        for pc in postcodes:
            idxs = [i for i, ln in enumerate(lines) if pc.replace(" ", "") in ln.replace(" ", "").upper()]
            for idx in idxs:
                start = max(0, idx - 3)
                end = min(len(lines), idx + 2)
                chunk = lines[start:end]
                addr = ", ".join([c.strip(" ,") for c in chunk if c.strip()])
                addr = re.sub(r"\s+", " ", addr).strip()

                score = 0.4
                if re.search(r"\b(road|rd|street|st|lane|ln|avenue|ave|drive|dr|close|cl|way|court|ct|flat|house)\b", addr, re.I):
                    score += 0.25
                if re.search(r"\b\d{1,4}\b", addr):
                    score += 0.15
                if len(addr) >= 20:
                    score += 0.10
                score = safe_clip01(score)

                if score > best[1]:
                    best = (addr, score, "postcode_window")

        if best[1] <= 0.0:
            return None, 0.0, "none"
        return best


# ---------------------------------------------------------------------
# Name extraction and reconciliation
# ---------------------------------------------------------------------
class NameResolver:
    def __init__(self, nlp):
        self.nlp = nlp

    @staticmethod
    def infer_name_from_filename(filename: str) -> Tuple[Optional[str], float]:
        stem = Path(filename).stem
        stem = re.sub(r"(?i)\b(resume|cv)\b", " ", stem)
        stem = re.sub(r"(?i)\b(final|latest|updated)\b", " ", stem)
        stem = re.sub(r"[_\.\-]+", " ", stem)
        stem = re.sub(r"([a-z])([A-Z])", r"\1 \2", stem)
        stem = re.sub(r"\b\d+\b", " ", stem)
        stem = re.sub(r"\s+", " ", stem).strip()

        if not stem:
            return None, 0.0

        tokens = stem.split()
        if len(tokens) < 2:
            return None, 0.0
        if len(tokens) > 6:
            tokens = tokens[:6]

        name = " ".join([t.capitalize() if not t.isupper() else t for t in tokens]).strip()
        if re.search(r"[^A-Za-z\s'’\-]", name):
            return None, 0.0

        conf = 0.65
        if 2 <= len(tokens) <= 4:
            conf = 0.80
        return name, conf

    def extract_name_from_content(self, text: str) -> Tuple[Optional[str], float]:
        if not text.strip():
            return None, 0.0

        lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
        header_lines = lines[:40]
        header_block = "\n".join(header_lines)

        bad_keywords = {
            "curriculum vitae", "cv", "resume", "profile", "summary",
            "contact", "details", "email", "e-mail", "phone", "mobile", "tel",
            "address", "postcode", "linkedin", "github", "nationality", "dob", "date of birth"
        }

        def strip_noise(s: str) -> str:
            s = re.sub(CONFIG.email_re, " ", s)
            s = re.sub(r"(\+44\s?\d[\d\s]{8,13}|\b0\d[\d\s]{8,13})", " ", s)
            s = re.sub(r"(?:https?://)?(?:www\.)?linkedin\.com/\S+", " ", s, flags=re.I)
            s = re.sub(r"(?:https?://)?(?:www\.)?github\.com/\S+", " ", s, flags=re.I)
            s = re.sub(r"\s+", " ", s).strip()
            return s

        def plausible_name(s: str) -> bool:
            if not s:
                return False
            if any(k in s.lower() for k in bad_keywords):
                return False
            if len(s) > 60:
                return False
            if re.search(r"[^A-Za-z\s'’\-\.]", s):
                return False
            parts = s.replace(".", "").split()
            if len(parts) < 2 or len(parts) > 5:
                return False
            if any(len(p) == 1 for p in parts) and sum(len(p) == 1 for p in parts) >= 2:
                return False
            return True

        candidates: List[Tuple[str, float]] = []
        name_line_re = re.compile(r"^([A-Za-z][A-Za-z'’\-]+(?:\s+[A-Za-z][A-Za-z'’\-]+){1,4})$")

        for i, ln in enumerate(header_lines[:15]):
            ln2 = strip_noise(ln)
            if not ln2:
                continue
            if any(k in ln2.lower() for k in bad_keywords):
                continue
            m = name_line_re.match(ln2)
            if m:
                nm = m.group(1).strip()
                score = 0.78 + max(0.0, (10 - i) * 0.01)
                candidates.append((nm, safe_clip01(score)))

            m2 = re.search(r"(?i)\bname\s*[:\-]\s*([A-Za-z][A-Za-z'’\-]+(?:\s+[A-Za-z][A-Za-z'’\-]+){1,4})\b", ln)
            if m2:
                nm = m2.group(1).strip()
                candidates.append((nm, 0.88))

        if self.nlp:
            try:
                doc = self.nlp(header_block)
                for ent in doc.ents:
                    if ent.label_ == "PERSON":
                        nm = strip_noise(ent.text).strip()
                        if nm.isupper():
                            nm = " ".join(w.capitalize() for w in nm.split())
                        if plausible_name(nm):
                            candidates.append((nm, 0.90))
            except Exception:
                pass

        if not candidates:
            return None, 0.0

        candidates.sort(key=lambda x: (x[1], len(x[0].split())), reverse=True)
        best = candidates[0]
        return best[0], best[1]

    def reconcile(self, content_name: Optional[str], content_conf: float, filename: str) -> Dict[str, Any]:
        file_name, file_conf = self.infer_name_from_filename(filename)

        if content_name and not file_name:
            return {"full_name": content_name, "name_source": "content", "name_confidence": safe_clip01(content_conf)}
        if file_name and not content_name:
            return {"full_name": file_name, "name_source": "filename", "name_confidence": safe_clip01(file_conf)}
        if not content_name and not file_name:
            return {"full_name": None, "name_source": "none", "name_confidence": 0.0}

        c = content_name or ""
        f = file_name or ""

        if RAPIDFUZZ_AVAILABLE:
            sim = fuzz.token_set_ratio(c, f) / 100.0
        else:
            sim = self._simple_token_similarity(c, f)

        if sim >= 0.90:
            conf = safe_clip01(max(content_conf, file_conf) + 0.05)
            return {"full_name": content_name, "name_source": "merged", "name_confidence": conf}
        if sim >= 0.75:
            merged = self._merge_names(c, f)
            conf = safe_clip01(max(content_conf, file_conf) - 0.05)
            return {"full_name": merged, "name_source": "merged", "name_confidence": conf}

        if content_conf >= file_conf:
            return {"full_name": content_name, "name_source": "content", "name_confidence": safe_clip01(content_conf - 0.10)}
        return {"full_name": file_name, "name_source": "filename", "name_confidence": safe_clip01(file_conf - 0.10)}

    @staticmethod
    def _simple_token_similarity(a: str, b: str) -> float:
        ta = set(a.lower().split())
        tb = set(b.lower().split())
        if not ta or not tb:
            return 0.0
        return len(ta & tb) / len(ta | tb)

    @staticmethod
    def _merge_names(a: str, b: str) -> str:
        ta = a.split()
        tb = b.split()
        out = []
        seen = set()
        for t in ta + tb:
            k = t.lower()
            if k not in seen:
                seen.add(k)
                out.append(t)
        return " ".join(out)


# ---------------------------------------------------------------------
# Licences and Skills and Experience and Keyword summariser
# (Your existing logic, unchanged where possible)
# ---------------------------------------------------------------------
class LicenceExtractor:
    def __init__(self, catalogue: SkillCatalogue):
        self.catalogue = catalogue
        self.keywords = [
            "sia", "door supervisor", "cctv", "first aid", "emergency first aid",
            "dbs", "cscs", "haccp", "food hygiene", "licence", "license", "certificate", "certification"
        ]

    def extract(self, sections: Dict[str, List[str]], full_text: str) -> List[Dict[str, Any]]:
        results: List[Dict[str, Any]] = []
        source_section = "certifications" if "certifications" in sections else None
        lines = sections.get("certifications", [])
        if not lines:
            lines = [ln.strip() for ln in full_text.split("\n") if ln.strip()]
            source_section = "content"

        for ln in lines:
            low = ln.lower()
            if len(ln) < 3:
                continue
            if any(k in low for k in self.keywords):
                canon, conf, method = self.catalogue.canonicalise(ln, context_text=full_text)
                results.append({
                    "text": ln,
                    "canonical": canon,
                    "confidence": safe_clip01(conf),
                    "source": source_section,
                    "method": method,
                })

        seen = set()
        out = []
        for r in results:
            key = (r["canonical"] or r["text"]).lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(r)
        return out


class SkillsExtractor:
    def __init__(self, nlp, catalogue: SkillCatalogue):
        self.nlp = nlp
        self.catalogue = catalogue
        self.stop_starts = {"responsible", "responsibilities", "duties", "role", "worked", "working"}

    def extract(self, sections: Dict[str, List[str]], full_text: str) -> Dict[str, Any]:
        mentions: List[Dict[str, Any]] = []

        skill_lines = sections.get("skills", [])
        mentions.extend(self._extract_from_lines(skill_lines, section="skills", context_text=full_text, base_conf=0.90))

        summary_lines = sections.get("summary", [])
        mentions.extend(self._extract_from_lines(summary_lines, section="summary", context_text=full_text, base_conf=0.70))

        exp_lines = sections.get("experience", [])
        exp_candidates = [ln for ln in exp_lines if CONFIG.bullet_re.match(ln) or (len(ln.split()) <= 12 and len(ln) >= 8)]
        mentions.extend(self._extract_from_lines(exp_candidates, section="experience", context_text=full_text, base_conf=0.65))

        mentions.extend(self._extract_dictionary_hits(full_text))

        canon_all = []
        for m in mentions:
            if m.get("canonical"):
                canon_all.append(m["canonical"])
        canon_all = dedupe_preserve_order(canon_all)

        if CONFIG.keep_unmatched_skill_phrases:
            for m in mentions:
                if m.get("canonical"):
                    continue
                if m.get("section") == "skills":
                    canon_all.append(m["text"])
            canon_all = dedupe_preserve_order(canon_all)

        seen = set()
        dedup_mentions = []
        for m in mentions:
            key = (m.get("canonical", "").lower(), m.get("text", "").lower(), m.get("section", "").lower())
            if key in seen:
                continue
            seen.add(key)
            dedup_mentions.append(m)

        dedup_mentions.sort(key=lambda x: x.get("confidence", 0.0), reverse=True)

        return {"all_skills": canon_all, "mentions": dedup_mentions}

    def _extract_from_lines(self, lines: List[str], section: str, context_text: str, base_conf: float) -> List[Dict[str, Any]]:
        out: List[Dict[str, Any]] = []
        for ln in lines:
            ln0 = ln.strip()
            if not ln0:
                continue
            ln0 = CONFIG.bullet_re.sub("", ln0).strip()

            candidates = self._split_skill_line(ln0)
            candidates.extend(self._phrase_mine(ln0))
            candidates = dedupe_preserve_order([c for c in candidates if len(c) >= 2])

            for c in candidates:
                canon, conf, method = self.catalogue.canonicalise(c, context_text=context_text)
                conf2 = conf
                if section == "skills":
                    conf2 = max(conf2, base_conf if method in {"exact", "variant", "abbrev"} else base_conf - 0.10)
                else:
                    conf2 = max(conf2, base_conf - 0.10)

                out.append({
                    "text": c,
                    "canonical": canon if method != "passthrough" else canon,
                    "section": section,
                    "sentence": ln.strip(),
                    "confidence": safe_clip01(conf2),
                    "method": method
                })
        return out

    def _extract_dictionary_hits(self, full_text: str) -> List[Dict[str, Any]]:
        out: List[Dict[str, Any]] = []
        low = full_text.lower()

        for canon in self.catalogue.canonical:
            if canon.lower() in low:
                out.append({"text": canon, "canonical": canon, "section": "content", "sentence": canon, "confidence": 0.85, "method": "dictionary"})

        for v, canon in self.catalogue.variant_to_canonical.items():
            if v in low:
                out.append({"text": v, "canonical": canon, "section": "content", "sentence": v, "confidence": 0.82, "method": "dictionary"})

        for abbr, canons in self.catalogue.abbreviations.items():
            if re.search(rf"\b{re.escape(abbr)}\b", full_text, flags=re.I):
                for c in canons[:1]:
                    out.append({"text": abbr, "canonical": c, "section": "content", "sentence": abbr, "confidence": 0.75, "method": "abbrev"})

        return out

    @staticmethod
    def _split_skill_line(line: str) -> List[str]:
        parts = [line.strip()]
        chunks = re.split(r"[;,/|]+", line)
        for ch in chunks:
            ch2 = ch.strip()
            if ch2 and ch2.lower() not in {"skills", "technical skills"}:
                parts.append(ch2)

        chunks2 = [c.strip() for c in re.split(r"\s*,\s*", line) if c.strip()]
        parts.extend(chunks2)

        if " and " in line.lower():
            and_parts = [p.strip() for p in re.split(r"(?i)\band\b", line) if p.strip()]
            if 2 <= len(and_parts) <= 4:
                parts.extend(and_parts)

        return [p for p in parts if len(p.split()) <= 14]

    def _phrase_mine(self, line: str) -> List[str]:
        out = []
        if not self.nlp:
            return out
        try:
            doc = self.nlp(line)
            for nc in getattr(doc, "noun_chunks", []):
                t = nc.text.strip().strip(" .,:;()[]{}")
                if 2 <= len(t.split()) <= 12:
                    out.append(t)
        except Exception:
            return out

        cleaned = []
        for t in out:
            t2 = re.sub(r"\s+", " ", t).strip()
            if not t2:
                continue
            if t2.lower().split()[0] in self.stop_starts:
                continue
            cleaned.append(t2)
        return cleaned


class ExperienceExtractor:
    def __init__(self, nlp, config: AppConfig):
        self.nlp = nlp
        self.config = config

    def extract(self, sections: Dict[str, List[str]], full_text: str) -> List[Dict[str, Any]]:
        exp_lines = sections.get("experience", [])
        if not exp_lines:
            exp_lines = [ln.strip() for ln in full_text.split("\n") if ln.strip()]

        blocks = self._split_into_role_blocks(exp_lines)
        roles = []
        for b in blocks:
            role = self._parse_role_block(b)
            if role:
                roles.append(role)
        return roles

    def _split_into_role_blocks(self, lines: List[str]) -> List[List[str]]:
        blocks: List[List[str]] = []
        buf: List[str] = []

        def flush():
            nonlocal buf
            if buf:
                cleaned = [x for x in buf if x.strip()]
                if cleaned:
                    blocks.append(cleaned)
            buf = []

        for ln in lines:
            if not ln.strip():
                if len(buf) >= 3:
                    flush()
                else:
                    buf.append(ln)
                continue

            if self.config.date_range_re.search(ln) and len(buf) >= 2:
                flush()
                buf.append(ln)
            else:
                buf.append(ln)

        flush()

        merged: List[List[str]] = []
        for b in blocks:
            if len(b) <= 2 and merged:
                merged[-1].extend(b)
            else:
                merged.append(b)
        return merged

    def _parse_role_block(self, lines: List[str]) -> Optional[Dict[str, Any]]:
        raw_block = [ln.strip() for ln in lines if ln.strip()]
        if len(raw_block) < 2:
            return None

        date_line_idx = None
        date_match = None
        for i, ln in enumerate(raw_block[:8]):
            m = self.config.date_range_re.search(ln)
            if m:
                date_line_idx = i
                date_match = m
                break

        header_lines = raw_block[:date_line_idx] if date_line_idx is not None else raw_block[:2]
        header_text = " | ".join(header_lines).strip()

        start_resp_idx = (date_line_idx + 1) if date_line_idx is not None else len(header_lines)
        resp_lines = raw_block[start_resp_idx:]
        responsibilities = self._clean_bullets(resp_lines)

        start_date, end_date, date_conf = self._parse_date_range(date_match.group(0)) if date_match else (None, None, 0.0)
        job_title, employer, location, header_conf = self._extract_header_fields(header_text)
        dur = self._duration_months(start_date, end_date)

        role = {
            "job_title": job_title,
            "employer": employer,
            "location": location,
            "start_date": start_date,
            "end_date": end_date,
            "duration_months": dur,
            "responsibilities": responsibilities,
            "confidence": {
                "job_title": safe_clip01(header_conf.get("job_title", 0.0)),
                "employer": safe_clip01(header_conf.get("employer", 0.0)),
                "location": safe_clip01(header_conf.get("location", 0.0)),
                "dates": safe_clip01(date_conf),
                "overall": safe_clip01(0.25 * header_conf.get("job_title", 0.0) + 0.25 * header_conf.get("employer", 0.0) + 0.20 * date_conf + 0.30)
            }
        }

        if not role["job_title"] and not role["employer"] and not date_match:
            return None
        return role

    @staticmethod
    def _clean_bullets(lines: List[str]) -> List[str]:
        out = []
        for ln in lines:
            t = ln.strip()
            if not t:
                continue
            t = CONFIG.bullet_re.sub("", t).strip()
            if looks_like_heading(t):
                continue
            if len(t) < 8:
                continue
            out.append(t)
        return dedupe_preserve_order(out)

    def _parse_date_range(self, s: str) -> Tuple[Optional[str], Optional[str], float]:
        m = self.config.date_range_re.search(s)
        if not m:
            return None, None, 0.0

        start_raw = m.group("start")
        end_raw = m.group("end")

        sd, sc = parse_month_year(start_raw)
        if sd is None:
            return None, None, 0.0

        end_lower = end_raw.strip().lower()
        if end_lower in {"present", "current", "now"}:
            return date_to_yyyy_mm(sd), "Current", 0.90

        ed, ec = parse_month_year(end_raw)
        if ed is None:
            return date_to_yyyy_mm(sd), None, 0.70

        conf = min(0.92, 0.60 + 0.20 * sc + 0.20 * ec)
        return date_to_yyyy_mm(sd), date_to_yyyy_mm(ed), conf

    def _extract_header_fields(self, header_text: str) -> Tuple[Optional[str], Optional[str], Optional[str], Dict[str, float]]:
        conf = {"job_title": 0.55, "employer": 0.45, "location": 0.35}

        parts = [p.strip() for p in re.split(r"\s*\|\s*|\s*-\s*|\s*•\s*|\s*,\s*", header_text) if p.strip()]
        parts = parts[:6]

        job_title = parts[0] if parts else None
        employer = parts[1] if len(parts) >= 2 else None
        location = None

        if self.nlp:
            try:
                doc = self.nlp(header_text)
                orgs = [ent.text.strip() for ent in doc.ents if ent.label_ in {"ORG"}]
                geos = [ent.text.strip() for ent in doc.ents if ent.label_ in {"GPE", "LOC"}]
                if orgs:
                    employer = orgs[0]
                    conf["employer"] = 0.75
                if geos:
                    location = geos[0]
                    conf["location"] = 0.70
            except Exception:
                pass

        if job_title and len(job_title.split()) <= 8:
            conf["job_title"] = 0.70
        if employer and len(employer.split()) <= 12:
            conf["employer"] = max(conf["employer"], 0.60)

        return job_title, employer, location, conf

    @staticmethod
    def _duration_months(start_date: Optional[str], end_date: Optional[str]) -> Optional[int]:
        if not start_date:
            return None
        try:
            sy, sm = [int(x) for x in start_date.split("-")]
            sd = date(sy, sm, 1)
        except Exception:
            return None

        if not end_date or end_date == "Current":
            return None

        try:
            ey, em = [int(x) for x in end_date.split("-")]
            ed = date(ey, em, 1)
        except Exception:
            return None

        if ed < sd:
            return None
        return months_between(sd, ed) + 1


class KeywordSummariser:
    def __init__(self):
        self.stop = {
            "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for",
            "of", "with", "by", "from", "as", "is", "was", "are", "were", "been",
            "this", "that", "these", "those", "i", "we", "they", "he", "she", "you",
        }

    def keywords(self, text: str, top_n: int = 35) -> List[str]:
        low = text.lower()
        words = re.findall(r"\b[a-z]{3,}\b", low)
        words = [w for w in words if w not in self.stop]
        if not words:
            return []
        freq = {}
        for w in words:
            freq[w] = freq.get(w, 0) + 1
        items = sorted(freq.items(), key=lambda x: x[1], reverse=True)
        return [k for k, _ in items[:top_n]]

    def summary(self, text: str, keywords: List[str], max_words: int = 160) -> str:
        lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
        scored = []
        top_kw = set([k.lower() for k in keywords[:18]])
        for ln in lines[:120]:
            if len(ln.split()) < 6:
                continue
            score = sum(1 for k in top_kw if k in ln.lower())
            if score > 0:
                scored.append((ln, score))
        scored.sort(key=lambda x: x[1], reverse=True)

        out = []
        wc = 0
        for ln, _ in scored[:6]:
            n = len(ln.split())
            if wc + n > max_words:
                continue
            out.append(ln)
            wc += n
        if out:
            return " ".join(out)
        return " ".join(text.split()[:max_words])


# ---------------------------------------------------------------------
# Main CV parser
# ---------------------------------------------------------------------
class CVParser:
    def __init__(self, nlp, catalogue: SkillCatalogue, config: AppConfig):
        self.config = config
        self.docx = DocumentExtractor()
        self.sections = SectionParser(config)
        self.contact = ContactExtractor(config)
        self.name = NameResolver(nlp)
        self.skills = SkillsExtractor(nlp, catalogue)
        self.licences = LicenceExtractor(catalogue)
        self.exp = ExperienceExtractor(nlp, config)
        self.kw = KeywordSummariser()

    def parse_bytes(self, file_bytes: bytes, filename: str, path_hint: str = "") -> Dict[str, Any]:
        t0 = time.time()
        extracted = self.docx.extract_text(file_bytes, filename)
        clean_text = extracted["clean_text"]

        sections = self.sections.split_sections(clean_text)

        content_name, content_conf = self.name.extract_name_from_content(clean_text)
        name_info = self.name.reconcile(content_name, content_conf, filename)

        emails = self.contact.extract_emails(clean_text)
        phones = self.contact.extract_uk_phones(clean_text)
        postcodes = self.contact.extract_postcodes(clean_text)
        addr_val, addr_conf, addr_src = self.contact.extract_address(clean_text, postcodes)

        skills = self.skills.extract(sections, clean_text)
        licences = self.licences.extract(sections, clean_text)
        experiences = self.exp.extract(sections, clean_text)

        keywords = self.kw.keywords(clean_text, top_n=35)
        summary = self.kw.summary(clean_text, keywords, max_words=160)

        elapsed = time.time() - t0

        result = {
            "metadata": {
                "filename": filename,
                "path": path_hint,
                "processed_at": datetime.now().isoformat(),
                "elapsed_seconds": round(elapsed, 3),
                "text_length": len(clean_text),
            },
            "personal_information": {
                "full_name": name_info["full_name"],
                "name_source": name_info["name_source"],
                "name_confidence": name_info["name_confidence"],
            },
            "contact_information": {
                "email": {"value": emails[0] if emails else None, "confidence": 0.95 if emails else 0.0, "source": "content"},
                "all_emails": emails,
                "phone": {"value": phones[0] if phones else None, "confidence": 0.90 if phones else 0.0, "source": "content"},
                "all_phones": phones,
                "postcode": {"value": postcodes[0] if postcodes else None, "confidence": 0.92 if postcodes else 0.0, "source": "content"},
                "all_postcodes": postcodes,
                "address": {"value": addr_val, "confidence": addr_conf, "source": addr_src},
            },
            "skills": skills,
            "professional_experience": experiences,
            "licences_and_certifications": licences,
            "keywords": keywords,
            "summary": summary,
        }
        return result


# ---------------------------------------------------------------------
# Sector storage and indexing
# ---------------------------------------------------------------------
def ensure_data_dir() -> Path:
    p = Path("data")
    p.mkdir(parents=True, exist_ok=True)
    return p


def load_sector_json(path: Path) -> Dict[str, Any]:
    if not path.exists():
        return {"sector": path.stem.replace("candidates", ""), "updated_at": None, "candidates": []}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {"sector": path.stem.replace("candidates", ""), "updated_at": None, "candidates": []}


def save_sector_json(path: Path, payload: Dict[str, Any]) -> None:
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def estimate_total_months_from_experience(entries: List[Dict[str, Any]]) -> int:
    total = 0
    for e in entries or []:
        m = e.get("duration_months")
        if isinstance(m, int) and m > 0:
            total += m
    return total


def build_candidate_search_text(parsed: Dict[str, Any]) -> str:
    pi = parsed.get("personal_information", {}) or {}
    ci = parsed.get("contact_information", {}) or {}
    skills = parsed.get("skills", {}) or {}
    exps = parsed.get("professional_experience", []) or []
    lic = parsed.get("licences_and_certifications", []) or []
    kw = parsed.get("keywords", []) or []

    parts = []
    parts.append(str(pi.get("full_name") or ""))
    parts.append(str(ci.get("email", {}).get("value") or ""))
    parts.append(str(ci.get("phone", {}).get("value") or ""))
    parts.append(str(ci.get("postcode", {}).get("value") or ""))
    parts.append(str(ci.get("address", {}).get("value") or ""))

    parts.extend([str(s) for s in (skills.get("all_skills") or [])])
    parts.extend([str(x.get("canonical") or x.get("text") or "") for x in lic])

    for e in exps:
        parts.append(str(e.get("job_title") or ""))
        parts.append(str(e.get("employer") or ""))
        parts.append(str(e.get("location") or ""))
        parts.extend([str(r) for r in (e.get("responsibilities") or [])])

    parts.extend([str(k) for k in kw])
    parts.append(str(parsed.get("summary") or ""))

    text = "\n".join([p for p in parts if p]).strip()
    return text


# ---------------------------------------------------------------------
# Query parsing and hybrid matching
# ---------------------------------------------------------------------
class QueryParser:
    years_re = re.compile(r"(?i)\b(\d{1,2})\s*\+?\s*(?:years|yrs|year)\b")
    postcode_re = CONFIG.uk_postcode_re

    def __init__(self, catalogue: SkillCatalogue):
        self.catalogue = catalogue

    def parse(self, q: str) -> Dict[str, Any]:
        q0 = (q or "").strip()
        low = q0.lower()

        years_min = None
        m = self.years_re.search(q0)
        if m:
            try:
                years_min = int(m.group(1))
            except Exception:
                years_min = None

        postcodes = []
        for m1, m2 in self.postcode_re.findall(q0):
            postcodes.append(normalise_postcode(m1, m2))

        raw_tokens = re.findall(r"[A-Za-z0-9][A-Za-z0-9\-\+\/& ]{1,40}", q0)
        raw_tokens = [t.strip() for t in raw_tokens if t.strip()]
        skill_hits = []
        for t in raw_tokens:
            canon, conf, method = self.catalogue.canonicalise(t, context_text=q0)
            if canon and conf >= 0.70 and method in {"exact", "variant", "abbrev", "embed"}:
                skill_hits.append(canon)
        skill_hits = dedupe_preserve_order(skill_hits)

        must_terms = []
        for w in re.findall(r"\b[a-z0-9]{2,}\b", low):
            if w in {"years", "yrs", "year", "experience", "exp"}:
                continue
            must_terms.append(w)
        must_terms = dedupe_preserve_order(must_terms)[:30]

        return {
            "raw": q0,
            "years_min": years_min,
            "postcodes": dedupe_preserve_order(postcodes),
            "skills": skill_hits,
            "must_terms": must_terms,
        }


class HybridMatcher:
    def __init__(self, nlp, catalogue: SkillCatalogue):
        self.nlp = nlp
        self.catalogue = catalogue
        self.query_parser = QueryParser(catalogue)

    def _tfidf_scores(self, query: str, docs: List[str]) -> List[float]:
        if not SKLEARN_AVAILABLE:
            return [0.0 for _ in docs]
        try:
            vec = TfidfVectorizer(
                lowercase=True,
                ngram_range=(1, 2),
                min_df=1,
                max_df=0.98,
                stop_words="english",
            )
            X = vec.fit_transform(docs + [query])
            qv = X[-1]
            dv = X[:-1]
            sims = (dv @ qv.T).toarray().reshape(-1)
            sims = np.clip(sims, 0.0, 1.0)
            return [float(x) for x in sims]
        except Exception:
            return [0.0 for _ in docs]

    def _embed_scores(self, query: str, docs: List[str]) -> List[float]:
        if not self.nlp:
            return [0.0 for _ in docs]
        try:
            qv = self.nlp(query).vector
            out = []
            for d in docs:
                dv = self.nlp(d[:5000]).vector
                out.append(safe_clip01((cosine_sim(qv, dv) + 1.0) / 2.0))
            return out
        except Exception:
            return [0.0 for _ in docs]

    @staticmethod
    def _keyword_overlap(query_terms: List[str], doc_text: str) -> float:
        if not query_terms:
            return 0.0
        low = doc_text.lower()
        hits = sum(1 for t in query_terms if t in low)
        return safe_clip01(hits / max(6, len(query_terms)))

    def rank(self, query: str, candidates: List[Dict[str, Any]], top_k: int = 25) -> List[Dict[str, Any]]:
        q = (query or "").strip()
        if not q:
            return []

        qp = self.query_parser.parse(q)

        docs = [c.get("_search_text", "") for c in candidates]
        tfidf = self._tfidf_scores(q, docs)
        emb = self._embed_scores(q, docs)

        ranked = []
        for i, c in enumerate(candidates):
            parsed = c.get("parsed") or {}
            ci = (parsed.get("contact_information") or {})
            si = (parsed.get("skills") or {})
            exps = parsed.get("professional_experience") or []

            postcode_val = (ci.get("postcode") or {}).get("value")
            total_months = c.get("total_experience_months", 0)
            total_years = total_months / 12.0 if total_months else 0.0

            if qp["postcodes"]:
                if not postcode_val or postcode_val.upper() not in [p.upper() for p in qp["postcodes"]]:
                    continue

            if qp["years_min"] is not None:
                if total_years + 0.01 < float(qp["years_min"]):
                    continue

            cand_skills = set([s.lower() for s in (si.get("all_skills") or [])])
            skill_bonus = 0.0
            if qp["skills"]:
                matched = sum(1 for s in qp["skills"] if s.lower() in cand_skills)
                skill_bonus = safe_clip01(matched / max(2, len(qp["skills"])))
            overlap = self._keyword_overlap(qp["must_terms"], c.get("_search_text", ""))

            # Hybrid score
            score = (
                0.45 * tfidf[i] +
                0.35 * emb[i] +
                0.10 * overlap +
                0.10 * skill_bonus
            )
            score = safe_clip01(score)
            print(score)

            ranked.append({
                "score": score,
                "candidate": c,
                "explain": {
                    "tfidf": tfidf[i],
                    "embed": emb[i],
                    "overlap": overlap,
                    "skill_bonus": skill_bonus,
                    "years": round(total_years, 2),
                    "postcode": postcode_val,
                    "query_skills": qp["skills"],
                }
            })

        ranked.sort(key=lambda x: x["score"], reverse=True)
        return ranked[:top_k]


# ---------------------------------------------------------------------
# Hardcoded sector folders
# You can paste WorkDrive folder links or folder IDs here.
# ---------------------------------------------------------------------
SECTOR_FOLDERS = {
    "security": {
        "folder": "https://workdrive.zoho.eu/3qg1t6735fa155c7a4eee8d68e1749c0f335e/privatespace/folders/hqzshe0f8092705f749378fc2dc6858d1e5e4",
        "json_file": "securitycandidates.json",
    },
    "warehouse": {
        "folder": "https://workdrive.zoho.eu/3qg1t6735fa155c7a4eee8d68e1749c0f335e/privatespace/folders/hqzsha77ea7cdf2614aea9869e053cbbfc73c",
        "json_file": "warehousecandidates.json",
    },
    "hospitality": {
        "folder": "https://workdrive.zoho.eu/3qg1t6735fa155c7a4eee8d68e1749c0f335e/privatespace/folders/c44zh82871ae5e9e44aafa2968f88ee3b9a3e",
        "json_file": "hospitalitycandidates.json",
    },
}


# ---------------------------------------------------------------------
# Sync pipeline
# ---------------------------------------------------------------------
def sync_sector(
    wd: ZohoWorkDriveClient,
    parser: CVParser,
    sector_key: str,
    include_subfolders: bool,
    exact_filename: str = "",
    show_debug: bool = False
) -> Dict[str, Any]:
    cfg = SECTOR_FOLDERS[sector_key]
    folder_id = wd.extract_resource_id(cfg["folder"])
    items = wd.list_resume_files_recursive(
        root_folder_id=folder_id,
        include_subfolders=include_subfolders,
        supported_exts=CONFIG.supported_exts,
        exact_filename=exact_filename.strip()
    )

    data_dir = ensure_data_dir()
    out_path = data_dir / cfg["json_file"]
    existing = load_sector_json(out_path)
    existing_map = { (x.get("file_id") or x.get("metadata", {}).get("filename") or ""): x for x in (existing.get("candidates") or []) }

    results = []
    errors = 0

    progress = st.progress(0)
    status = st.empty()

    for i, it in enumerate(items, start=1):
        status.write(f"Syncing {sector_key}. {i}/{len(items)}. {it.get('virtual_path')}")
        try:
            file_id = it["id"]
            filename = it["name"]
            web_url = wd.file_web_url(file_id)

            b = wd.download_file_bytes(file_id)
            parsed = parser.parse_bytes(file_bytes=b, filename=filename, path_hint=f"workdrive:{it.get('virtual_path')}")
            total_months = estimate_total_months_from_experience(parsed.get("professional_experience") or [])
            search_text = build_candidate_search_text(parsed)

            candidate_record = {
                "sector": sector_key,
                "file_id": file_id,
                "file_name": filename,
                "file_path": it.get("virtual_path"),
                "file_url": web_url,
                "synced_at": datetime.now().isoformat(),
                "total_experience_months": int(total_months),
                "parsed": parsed,
                "_search_text": search_text,
            }
            existing_map[file_id] = candidate_record

        except Exception as ex:
            errors += 1
            if show_debug:
                st.code(traceback.format_exc())
            continue

        progress.progress(int(i * 100 / max(1, len(items))))

    results = list(existing_map.values())
    results.sort(key=lambda x: (x.get("file_name") or "").lower())

    payload = {
        "sector": sector_key,
        "updated_at": datetime.now().isoformat(),
        "total_files_seen": len(items),
        "total_candidates": len(results),
        "errors": errors,
        "candidates": results
    }
    save_sector_json(out_path, payload)
    return payload


def load_all_candidates() -> List[Dict[str, Any]]:
    data_dir = ensure_data_dir()
    all_cands = []
    for sector_key, cfg in SECTOR_FOLDERS.items():
        p = data_dir / cfg["json_file"]
        blob = load_sector_json(p)
        for c in (blob.get("candidates") or []):
            all_cands.append(c)
    return all_cands


# ---------------------------------------------------------------------
# Streamlit UI
# ---------------------------------------------------------------------
def main():
    st.set_page_config(page_title="WorkDrive CV Sync + Search", page_icon="📄", layout="wide")
    st.title("📄 WorkDrive CV Sync + Advanced Search")
    st.caption("Three sectors. One sync button per sector. Hybrid NLP matching for fast shortlisting.")

    nlp = load_spacy_model()
    if not nlp:
        st.warning("spaCy model not loaded. Search will still work, but semantic matching will be weaker.")

    with st.sidebar:
        st.header("⚙️ Settings")
        include_subfolders = st.checkbox("Search all subfolders", value=True)
        exact_filename = st.text_input("Exact filename (optional)", value="")
        show_debug = st.checkbox("Show debug logs", value=False)

        st.markdown("---")
        enable_ocr = st.checkbox("Enable OCR fallback for scanned PDFs", value=CONFIG.enable_ocr_fallback)
        CONFIG.enable_ocr_fallback = enable_ocr

        sim_threshold = st.slider(
            "Embedding similarity threshold",
            min_value=0.55,
            max_value=0.90,
            value=float(CONFIG.embedding_similarity_threshold),
            step=0.01
        )
        CONFIG.embedding_similarity_threshold = float(sim_threshold)

        st.markdown("---")
        st.write("Storage")
        st.code("data/securitycandidates.json\ndata/warehousecandidates.json\ndata/hospitalitycandidates.json", language="text")

    catalogue = SkillCatalogue.load_from_path(None, nlp=nlp)
    parser = CVParser(nlp, catalogue, CONFIG)
    matcher = HybridMatcher(nlp, catalogue)

    # Sync area
    st.subheader("1) Sync files from WorkDrive into sector JSON")
    cols = st.columns(3)

    try:
        wd = ZohoWorkDriveClient()
    except Exception as e:
        st.error(str(e))
        st.stop()

    for idx, sector_key in enumerate(["security", "warehouse", "hospitality"]):
        with cols[idx]:
            st.markdown(f"### {sector_key.capitalize()}")
            st.caption(f"Folder: {SECTOR_FOLDERS[sector_key]['folder']}")
            if st.button(f"🔄 Sync {sector_key}", type="primary", use_container_width=True):
                try:
                    if "PASTE_" in SECTOR_FOLDERS[sector_key]["folder"]:
                        st.error(f"Update SECTOR_FOLDERS['{sector_key}']['folder'] with your WorkDrive folder link or ID.")
                    else:
                        payload = sync_sector(
                            wd=wd,
                            parser=parser,
                            sector_key=sector_key,
                            include_subfolders=include_subfolders,
                            exact_filename=exact_filename.strip(),
                            show_debug=show_debug
                        )
                        st.success(f"Synced {payload.get('total_candidates')} candidates. Errors: {payload.get('errors')}")
                except Exception as e:
                    st.error(str(e))
                    if show_debug:
                        st.code(traceback.format_exc())

    st.markdown("---")

    # Search area
    st.subheader("2) Search candidates")
    all_candidates = load_all_candidates()
    st.caption(f"Loaded candidates in local JSON: {len(all_candidates)}")

    q = st.text_input(
        "Search query",
        value="",
        placeholder='Examples: "SIA Door Supervisor 5 years TW3 3NU" . "forklift reach driver slough" . "housekeeping room attendant"',
    )
    

    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        sector_filter = st.multiselect("Filter sectors", options=["security", "warehouse", "hospitality"], default=["security", "warehouse", "hospitality"])
    with c2:
        top_k = st.slider("Top results", min_value=5, max_value=50, value=20, step=5)
    with c3:
        min_score = st.slider("Minimum score", min_value=0.0, max_value=1.0, value=0.25, step=0.05)

    filtered = [c for c in all_candidates if c.get("sector") in set(sector_filter)]

    # Placeholders to control UI
    results_box = st.container()          # where results will be shown
    details_box = st.container()          # where expanders will be shown
    status_box = st.empty()              # small text status


    if q.strip():
        # clear previous UI first
        results_box.empty()
        details_box.empty()
        status_box.empty()

        with results_box:
            with st.spinner("Searching candidates..."):
                t0 = time.perf_counter()
                ranked = matcher.rank(q, filtered, top_k=int(top_k))
                ranked = [r for r in ranked if r["score"] >= float(min_score)]
                dt = time.perf_counter() - t0

        status_box.caption(f"Search completed in {dt:.2f} seconds. Results: {len(ranked)}")

        if not ranked:
            results_box.info("No matches found with current filters.")
        else:
            rows = []
            for r in ranked:
                c = r["candidate"]
                parsed = c.get("parsed") or {}
                pi = parsed.get("personal_information") or {}
                ci = parsed.get("contact_information") or {}
                si = parsed.get("skills") or {}

                rows.append({
                    "score": round(r["score"], 4),
                "sector": c.get("sector"),
                "name": pi.get("full_name"),
                "email": (ci.get("email") or {}).get("value"),
                "phone": (ci.get("phone") or {}).get("value"),
                "postcode": (ci.get("postcode") or {}).get("value"),
                "years_est": round((c.get("total_experience_months", 0) / 12.0), 2),
                "top_skills": ", ".join((si.get("all_skills") or [])[:12]),
                "file": c.get("file_name"),
                "resume_url": c.get("file_url"),
            })

            df = pd.DataFrame(rows)

            with results_box:
                st.dataframe(df, use_container_width=True)

            with details_box:
                st.markdown("---")
                st.subheader("Top matches details")
                for i, r in enumerate(ranked[:min(len(ranked), 10)], start=1):
                    c = r["candidate"]
                    parsed = c.get("parsed") or {}
                    pi = parsed.get("personal_information") or {}
                    ci = parsed.get("contact_information") or {}
                    si = parsed.get("skills") or {}

                    title = f"{i}. {pi.get('full_name') or c.get('file_name')} . Score {round(r['score'], 3)} . {c.get('sector')}"
                    with st.expander(title, expanded=(i == 1)):
                        st.markdown(f"**Resume:** {c.get('file_url')}")
                        st.write("**Email:**", (ci.get("email") or {}).get("value"))
                        st.write("**Phone:**", (ci.get("phone") or {}).get("value"))
                        st.write("**Postcode:**", (ci.get("postcode") or {}).get("value"))
                        st.write("**Estimated experience (years):**", round((c.get("total_experience_months", 0) / 12.0), 2))
                        st.write("**Top skills:**", ", ".join((si.get("all_skills") or [])[:25]))
                        st.write("**Summary:**")
                        st.info(parsed.get("summary") or "")

                        if show_debug:
                            st.write("Explain")
                            st.json(r.get("explain") or {})
    else:
        results_box.empty()
        details_box.empty()
        status_box.info("Type a query to rank candidates. Example: SIA Door Supervisor 5 years TW3 3NU")




if __name__ == "__main__":
    main()
